from pathlib import Path
from docx import Document

paths = [
    Path('Bulletpont questions.docx'),
    Path('Bulletpont Memo.docx'),
    Path('Bulletpont learner data.docx'),
    Path('Paragraph formatting Questions.docx'),
    Path('Paragraph formatting Memo.docx'),
    Path('Paragraph formatting Learner DATA.docx'),
    Path('COMPUTER APPLICATIONS TECHNOLOGY test 1 qp.docx'),
    Path('COMPUTER APPLICATIONS TECHNOLOGY test 1 memo.docx'),
    Path('COMPUTER APPLICATIONS TECHNOLOGY test 1 (3).docx'),
]

for p in paths:
    print('='*80)
    print('FILE:', p)
    doc = Document(p)
    print('paragraphs:', len(doc.paragraphs), 'tables:', len(doc.tables))
    for i, para in enumerate(doc.paragraphs[:20], 1):
        text = para.text.strip()
        if text:
            print(f'P{i}:', text[:200])
    for ti, table in enumerate(doc.tables, 1):
        print(f'TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}')
        for ri, row in enumerate(table.rows[:10], 1):
            print(' ', ri, [cell.text.strip()[:100] for cell in row.cells])
    print()
