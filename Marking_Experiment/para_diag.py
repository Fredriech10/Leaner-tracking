import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).parent.parent
_align = {0: 'left', 1: 'center', 2: 'right', 3: 'justify', None: 'None'}

files = [
    ROOT / 'Paragraph formatting Questions.docx',
    ROOT / 'Paragraph formatting Memo.docx',
    ROOT / 'Paragraph formatting Learner DATA.docx',
]

out = []

for fpath in files:
    out.append(f'\n======== {fpath.name} ========')
    doc = Document(fpath)
    out.append(f'  Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}')
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        fmt = p.paragraph_format
        try:
            ls = round(float(fmt.line_spacing), 2) if fmt.line_spacing else None
        except Exception:
            ls = fmt.line_spacing
        try:
            sb = round(fmt.space_before.pt, 1) if fmt.space_before else None
        except Exception:
            sb = None
        try:
            sa = round(fmt.space_after.pt, 1) if fmt.space_after else None
        except Exception:
            sa = None
        al = _align.get(p.alignment, str(p.alignment))
        run_info = ''
        for r in p.runs:
            if r.text.strip():
                try:
                    sz = round(r.font.size.pt, 1) if r.font.size else None
                except Exception:
                    sz = None
                run_info = f'font={r.font.name} sz={sz} bold={r.font.bold} italic={r.font.italic}'
                break
        out.append(f'  [P{i}] style={p.style.name!r} al={al} ls={ls} sb={sb} sa={sa} | {txt[:70]} | {run_info}')

    for ti, t in enumerate(doc.tables):
        out.append(f'  [T{ti}]:')
        for ri, row in enumerate(t.rows):
            cells = [c.text.strip()[:50] for c in row.cells]
            out.append(f'    row{ri}: {cells}')

output_path = ROOT / 'Marking_Experiment' / 'para_diag.txt'
output_path.write_text('\n'.join(out), encoding='utf-8')
print(f'Written to {output_path}')
