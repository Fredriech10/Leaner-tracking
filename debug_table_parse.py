from pathlib import Path
from Marking_Experiment.marksheet_parser import MarksheetParser
from docx import Document

memo = Path(r'd:\Sripts\Leaner tracking\COMPUTER APPLICATIONS TECHNOLOGY test 1 memo.docx')

doc = Document(memo)
paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
print('paras', len(paras))
print('first paras', paras[:10])

tables = []
for i, table in enumerate(doc.tables, 1):
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    tables.append((i, rows))
    print('TABLE', i, 'rows', len(rows), 'cols max', max(len(r) for r in rows))
    for r in rows[:15]:
        print(r)

text = '\n\n'.join(paras + [f'TABLE {i}:\n' + '\n'.join(['\t'.join(row) for row in rows]) for i, rows in tables])
print('\n--- text preview ---\n')
print(text[:1600])

parser = MarksheetParser(backend='ollama')
blocks = parser._extract_table_blocks(text)
print('blocks', len(blocks))
for bi, block in enumerate(blocks, 1):
    print('BLOCK', bi, 'rows', len(block))
    for r in block[:15]:
        print(r)
