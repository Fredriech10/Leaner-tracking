"""Table rule checker for Word documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document

from ..checker_types import CheckerResult
from .word_utils import _find_tables, _find_paragraphs


def check_table_rule(rule: Dict[str, Any], file_path: Path) -> CheckerResult:
    document = Document(file_path)
    check_type = str(rule.get("type", ""))
    target = rule.get("target", {}) or {}
    expected = rule.get("expected")

    tables = _find_tables(document, target)
    if not tables:
        return CheckerResult(passed=False, details={"reason": "Table target not found."})
    table = tables[0]
    actual = None
    passed = False

    if check_type == "exists":
        passed = True
        actual = f"Table with {len(table.rows)} rows, {len(table.columns)} columns"

    elif check_type == "row_count":
        actual = len(table.rows)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "column_count":
        actual = len(table.columns)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "cell_content":
        cell_target = target.get("cell", {})
        row_idx = cell_target.get("row", 0)
        col_idx = cell_target.get("column", 0)
        if row_idx < len(table.rows) and col_idx < len(table.columns):
            cell = table.cell(row_idx, col_idx)
            actual = cell.text.strip()
            if isinstance(expected, dict):
                exp_text = expected.get("text", "")
                exp_case_sensitive = expected.get("case_sensitive", False)
                if not exp_case_sensitive:
                    passed = actual.lower() == exp_text.lower()
                else:
                    passed = actual == exp_text
            else:
                passed = actual == str(expected)
        else:
            passed = False
            actual = "Cell out of range"

    elif check_type == "cell_formatting":
        cell_target = target.get("cell", {})
        row_idx = cell_target.get("row", 0)
        col_idx = cell_target.get("column", 0)
        if row_idx < len(table.rows) and col_idx < len(table.columns):
            cell = table.cell(row_idx, col_idx)
            passed, actual = _check_cell_formatting(cell, expected)
        else:
            passed = False
            actual = "Cell out of range"

    elif check_type == "border":
        passed, actual = _check_table_border(table, expected, file_path)

    elif check_type == "shading":
        passed, actual = _check_table_shading(table, expected)

    elif check_type == "alignment":
        passed, actual = _check_table_alignment(table, expected)

    else:
        return CheckerResult(passed=False, details={"reason": "Unsupported table check."})

    return CheckerResult(passed=passed, actual=actual, details={"type": check_type})


def _check_cell_formatting(cell, expected: Any) -> tuple[bool, Any]:
    """Check formatting of a table cell."""
    actual = {}
    passed = True
    if isinstance(expected, dict):
        if "bold" in expected:
            bold_found = any(run.font.bold for para in cell.paragraphs for run in para.runs)
            actual["bold"] = bold_found
            if bold_found != bool(expected["bold"]):
                passed = False
        if "italic" in expected:
            italic_found = any(run.font.italic for para in cell.paragraphs for run in para.runs)
            actual["italic"] = italic_found
            if italic_found != bool(expected["italic"]):
                passed = False
        if "color" in expected:
            # Simplified color check - check first run
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.color.rgb:
                        actual["color"] = run.font.color.rgb
                        # Basic color match - could be enhanced
                        if str(run.font.color.rgb).lower() != str(expected["color"]).lower():
                            passed = False
                        break
                if "color" in actual:
                    break
    return passed, actual


def _check_table_border(table, expected: Any, file_path: Path) -> tuple[bool, Any]:
    """Check table border via raw XML."""
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        tbl_xml = etree.fromstring(table._tbl.xml.encode("utf-8"))
    except Exception:
        return False, None
    tblBdr = tbl_xml.find(f".//{{{WNS}}}tblBdr")
    if tblBdr is None:
        return False, {"reason": "No table border found"}
    sides = ["top", "bottom", "left", "right", "insideH", "insideV"]
    actual = {}
    for side in sides:
        el = tblBdr.find(f"{{{WNS}}}{side}")
        if el is not None:
            actual[side] = {
                "style": el.get(f"{{{WNS}}}val", ""),
                "color": el.get(f"{{{WNS}}}color", "").lower(),
                "sz": el.get(f"{{{WNS}}}sz", ""),
            }
    if not actual:
        return False, {"reason": "No border sides found"}
    passed = True
    if isinstance(expected, dict):
        exp_style = expected.get("style")
        exp_color = expected.get("color", "").lower() if expected.get("color") else None
        for side_data in actual.values():
            if exp_style and exp_style.lower() not in side_data["style"].lower():
                passed = False
            if exp_color and side_data["color"] not in [exp_color]:
                passed = False
    return passed, actual


def _check_table_shading(table, expected: Any) -> tuple[bool, Any]:
    """Check table shading via raw XML."""
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        tbl_xml = etree.fromstring(table._tbl.xml.encode("utf-8"))
    except Exception:
        return False, None
    shd = tbl_xml.find(f".//{{{WNS}}}shd")
    if shd is None:
        return False, {"reason": "No shading found"}
    fill = shd.get(f"{{{WNS}}}fill", "").lower()
    actual = {"fill": fill}
    if isinstance(expected, dict):
        exp_color = expected.get("color", "any").lower()
        if exp_color == "any":
            passed = bool(fill) and fill != "auto"
        else:
            passed = fill == exp_color
    else:
        passed = bool(fill) and fill != "auto"
    return passed, actual


def _check_table_alignment(table, expected: Any) -> tuple[bool, Any]:
    """Check table alignment."""
    # Table alignment is tricky - check first cell or overall
    try:
        alignment = table.alignment
        actual = str(alignment) if alignment else "left"
        passed = actual.lower() == str(expected).lower()
    except Exception:
        passed = False
        actual = "unknown"
    return passed, actual