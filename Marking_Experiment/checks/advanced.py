"""Advanced rule checker for Word documents (complex formatting, styles, etc.)."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document

from checker_types import CheckerResult
from .word_utils import _find_paragraphs


def check_advanced_rule(rule: Dict[str, Any], file_path: Path) -> CheckerResult:
    document = Document(file_path)
    check_type = str(rule.get("type", ""))
    target = rule.get("target", {}) or {}
    expected = rule.get("expected")

    actual = None
    passed = False

    if check_type == "style_applied":
        paragraphs = _find_paragraphs(document, target)
        if paragraphs:
            style_names = [p.style.name for p in paragraphs if p.style]
            actual = style_names
            if isinstance(expected, list):
                passed = all(style in style_names for style in expected)
            else:
                passed = str(expected) in style_names
        else:
            passed = False
            actual = "No paragraphs found"

    elif check_type == "custom_style":
        # Check if custom styles are used
        custom_styles = [s for s in document.styles if not s.builtin]
        actual = len(custom_styles)
        try:
            min_count = expected.get("min", 1) if isinstance(expected, dict) else int(expected)
            passed = actual >= min_count
        except Exception:
            passed = actual > 0

    elif check_type == "section_breaks":
        # Count section breaks
        actual = len(document.sections)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "columns":
        # Check if document uses columns
        section = document.sections[0]
        actual = section.columns.num
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "watermark":
        # Check for watermark
        watermark_found = _find_watermark(document)
        actual = "Found" if watermark_found else "Not found"
        passed = watermark_found == bool(expected)

    elif check_type == "track_changes":
        # Check if track changes is enabled
        track_changes = _check_track_changes(document)
        actual = "Enabled" if track_changes else "Disabled"
        passed = track_changes == bool(expected)

    elif check_type == "comments":
        # Count comments
        comments = _find_comments(document)
        actual = len(comments)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "revision_count":
        # Count revisions
        revisions = _count_revisions(document)
        actual = revisions
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "field_codes":
        # Check for field codes
        fields = _find_field_codes(document)
        actual = len(fields)
        try:
            min_count = expected.get("min", 1) if isinstance(expected, dict) else int(expected)
            passed = actual >= min_count
        except Exception:
            passed = actual > 0

    elif check_type == "macros":
        # Check for macros (VBA)
        macros_found = _check_macros(document)
        actual = "Found" if macros_found else "Not found"
        passed = macros_found == bool(expected)

    elif check_type == "password_protected":
        # Check if document is password protected
        protected = _check_password_protection(document)
        actual = "Protected" if protected else "Not protected"
        passed = protected == bool(expected)

    elif check_type == "digital_signature":
        # Check for digital signatures
        signed = _check_digital_signature(document)
        actual = "Signed" if signed else "Not signed"
        passed = signed == bool(expected)

    else:
        return CheckerResult(passed=False, details={"reason": "Unsupported advanced check."})

    return CheckerResult(passed=passed, actual=actual, details={"type": check_type})


def _find_watermark(document) -> bool:
    """Check if document has a watermark."""
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for section in document.sections:
            header = section.header
            if header:
                h_xml = etree.fromstring(header._hdr.xml.encode("utf-8"))
                watermark = h_xml.find(f".//{{{WNS}}}pict")
                if watermark is not None:
                    return True
    except Exception:
        pass
    return False


def _check_track_changes(document) -> bool:
    """Check if track changes is enabled."""
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        settings = document.settings
        s_xml = etree.fromstring(settings._settings.xml.encode("utf-8"))
        track_revisions = s_xml.find(f".//{{{WNS}}}trackRevisions")
        if track_revisions is not None:
            val = track_revisions.get(f"{{{WNS}}}val", "false")
            return val.lower() == "true"
    except Exception:
        pass
    return False


def _find_comments(document) -> list[Dict[str, Any]]:
    """Find all comments in document."""
    comments = []
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            comment_refs = p_xml.findall(f".//{{{WNS}}}commentReference")
            for ref in comment_refs:
                comments.append({"id": ref.get("{{{WNS}}}id", "")})
    except Exception:
        pass
    return comments


def _count_revisions(document) -> int:
    """Count revisions in document."""
    count = 0
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            # Look for revision marks
            ins = p_xml.findall(f".//{{{WNS}}}ins")
            del_ = p_xml.findall(f".//{{{WNS}}}del")
            count += len(ins) + len(del_)
    except Exception:
        pass
    return count


def _find_field_codes(document) -> list[Dict[str, Any]]:
    """Find all field codes in document."""
    fields = []
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            fld_simple = p_xml.findall(f".//{{{WNS}}}fldSimple")
            for field in fld_simple:
                instr = field.get(f"{{{WNS}}}instr", "")
                fields.append({"instruction": instr})
    except Exception:
        pass
    return fields


def _check_macros(document) -> bool:
    """Check if document contains macros."""
    # This is a simplified check - macros are stored in VBA project
    try:
        # Check if there's a VBA project relationship
        for rel in document.part.rels.values():
            if "vbaProject" in rel.reltype:
                return True
    except Exception:
        pass
    return False


def _check_password_protection(document) -> bool:
    """Check if document is password protected."""
    # Password protection info is in document settings
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        settings = document.settings
        s_xml = etree.fromstring(settings._settings.xml.encode("utf-8"))
        protection = s_xml.find(f".//{{{WNS}}}documentProtection")
        if protection is not None:
            return True
    except Exception:
        pass
    return False


def _check_digital_signature(document) -> bool:
    """Check if document has digital signatures."""
    try:
        # Digital signatures are in the package relationships
        for rel in document.part.rels.values():
            if "digital-signature" in rel.reltype.lower():
                return True
    except Exception:
        pass
    return False