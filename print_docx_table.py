from pathlib import Path
from docx import Document

for name in [
    'Bulletpont Memo.docx',
    'Paragraph formatting Memo.docx',
    'COMPUTER APPLICATIONS TECHNOLOGY test 1 memo.docx',
]:
    path = Path(name)
    print('='*80)
    print('FILE:', path)
    if not path.exists():
        print('MISSING')
        continue
    doc = Document(path)
    print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables))
    for ti, table in enumerate(doc.tables, start=1):
        print(f'TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}')
        for ri, row in enumerate(table.rows, start=1):
            print(' ', ri, [cell.text.strip() for cell in row.cells])
    print()
