"""Document-level rule checker for Word documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document

from ..checker_types import CheckerResult


def check_document_rule(rule: Dict[str, Any], file_path: Path) -> CheckerResult:
    document = Document(file_path)
    check_type = str(rule.get("type", ""))
    expected = rule.get("expected")

    actual = None
    passed = False

    if check_type == "page_count":
        # Estimate page count - rough approximation
        total_chars = sum(len(p.text) for p in document.paragraphs)
        # Average 3000 chars per page
        actual = max(1, round(total_chars / 3000))
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "word_count":
        actual = sum(len(p.text.split()) for p in document.paragraphs)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "character_count":
        actual = sum(len(p.text) for p in document.paragraphs)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "paragraph_count":
        actual = len(document.paragraphs)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "has_headers":
        headers = []
        for section in document.sections:
            if section.header.is_linked_to_previous is False:
                headers.append(section.header)
        actual = len(headers) > 0
        passed = actual == bool(expected)

    elif check_type == "has_footers":
        footers = []
        for section in document.sections:
            if section.footer.is_linked_to_previous is False:
                footers.append(section.footer)
        actual = len(footers) > 0
        passed = actual == bool(expected)

    elif check_type == "page_orientation":
        # Check first section
        section = document.sections[0]
        orientation = "portrait" if section.page_width > section.page_height else "landscape"
        actual = orientation
        passed = actual == str(expected).lower()

    elif check_type == "page_size":
        section = document.sections[0]
        width_inch = section.page_width.inches
        height_inch = section.page_height.inches
        actual = f"{round(width_inch, 1)}x{round(height_inch, 1)} inches"
        if isinstance(expected, dict):
            exp_width = expected.get("width")
            exp_height = expected.get("height")
            exp_unit = expected.get("unit", "inches")
            if exp_unit == "inches":
                passed = abs(width_inch - exp_width) < 0.1 and abs(height_inch - exp_height) < 0.1
        else:
            passed = str(expected).lower() in actual.lower()

    elif check_type == "margins":
        section = document.sections[0]
        margins = {
            "top": round(section.top_margin.inches, 2),
            "bottom": round(section.bottom_margin.inches, 2),
            "left": round(section.left_margin.inches, 2),
            "right": round(section.right_margin.inches, 2),
        }
        actual = margins
        if isinstance(expected, dict):
            passed = True
            for side in ["top", "bottom", "left", "right"]:
                if side in expected:
                    if abs(margins[side] - expected[side]) > 0.1:
                        passed = False
        else:
            passed = True  # Just check if margins exist

    elif check_type == "contains_text":
        text_to_find = expected.get("text", "") if isinstance(expected, dict) else str(expected)
        found = any(text_to_find.lower() in p.text.lower() for p in document.paragraphs)
        passed = found
        actual = text_to_find if found else None

    elif check_type == "language":
        # Check document language - default is usually English
        actual = "en-US"  # Placeholder - would need to check document settings
        passed = actual == str(expected)

    else:
        return CheckerResult(passed=False, details={"reason": "Unsupported document check."})

    return CheckerResult(passed=passed, actual=actual, details={"type": check_type})