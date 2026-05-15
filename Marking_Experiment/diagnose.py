"""Diagnostic for any set of docx files."""
import sys
from pathlib import Path
from zipfile import ZipFile
import re

ROOT = Path(__file__).parent.parent

files = [
    ROOT / 'Bulletpont questions.docx',
    ROOT / 'Bulletpont Memo.docx',
    ROOT / 'Bulletpont learner data.docx',
]

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
_align = {0:'left',1:'center',2:'right',3:'justify',None:'None'}

out = []

for fpath in files:
    out.append(f'\n======== {fpath.name} ========')
    doc = Document(fpath)
    out.append(f'  Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}')
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        fmt = p.paragraph_format
        try: ls = round(float(fmt.line_spacing),2) if fmt.line_spacing else None
        except: ls = fmt.line_spacing
        ls_rule = fmt.line_spacing_rule
        try: sb = round(fmt.space_before.pt,1) if fmt.space_before else None
        except: sb = None
        try: sa = round(fmt.space_after.pt,1) if fmt.space_after else None
        except: sa = None
        fi = fmt.first_line_indent
        fi_cm = round(fi/360000.0,3) if fi else None
        al = _align.get(p.alignment, str(p.alignment))
        # numbering
        numPr = None
        if p._p.pPr is not None:
            numPr = p._p.pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        has_num = numPr is not None
        run_info = ''
        for r in p.runs:
            if r.text.strip():
                try: sz = round(r.font.size.pt,1) if r.font.size else None
                except: sz = None
                run_info = f'font={r.font.name} sz={sz} bold={r.font.bold} italic={r.font.italic}'
                break
        out.append(f'  [P{i}] style={p.style.name!r} al={al} ls={ls} ls_rule={ls_rule} sb={sb} sa={sa} fi_cm={fi_cm} num={has_num} | {txt[:60]} | {run_info}')
    for ti, t in enumerate(doc.tables):
        out.append(f'  [T{ti}]:')
        for ri, row in enumerate(t.rows):
            cells = [c.text.strip()[:50] for c in row.cells]
            out.append(f'    row{ri}: {cells}')

output_path = ROOT / 'Marking_Experiment' / 'bullet_diag.txt'
output_path.write_text('\n'.join(out), encoding='utf-8')
print(f'Written to {output_path}')

