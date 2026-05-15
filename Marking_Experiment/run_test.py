"""Standalone test runner — no LLM required, uses heuristic fallback."""

from __future__ import annotations

import json
import sys
from pathlib import Path

# Allow running as: python -m Marking_Experiment.run_test  OR  python run_test.py
ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

from Marking_Experiment.marksheet_parser import MarksheetParser
from Marking_Experiment.engine import MarkingEngine

QP_PATH    = ROOT / "Bulletpont questions.docx"
MEMO_PATH  = ROOT / "Bulletpont Memo.docx"
SUBMISSION = ROOT / "Bulletpont learner data.docx"


def read_docx_rich(path: Path) -> str:
    """Extract text + formatting hints from a docx for the parser."""
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    _align = {
        WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
        WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
        WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
        WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
    }
    doc = Document(path)
    parts = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        hints = []
        align = _align.get(p.alignment)
        if align:
            hints.append(f"align={align}")
        fmt = p.paragraph_format
        if fmt.line_spacing is not None:
            hints.append(f"line_spacing={fmt.line_spacing}")
        for run in p.runs:
            if run.text.strip():
                if run.font.size:
                    try:
                        hints.append(f"font_size={run.font.size.pt}pt")
                    except Exception:
                        pass
                if run.font.bold:
                    hints.append("bold")
                if run.font.italic:
                    hints.append("italic")
                break
        style = p.style.name if p.style else ""
        hint_str = f" [{', '.join(hints)}]" if hints else ""
        style_str = f" (style: {style})" if style else ""
        parts.append(f"[P{idx}]{style_str} {text}{hint_str}")
    for t_idx, table in enumerate(doc.tables, 1):
        rows = ["\t".join(c.text.strip() for c in row.cells) for row in table.rows]
        parts.append(f"TABLE {t_idx}:\n" + "\n".join(rows))
    return "\n\n".join(parts)


def main() -> None:
    print("=" * 60)
    print("MARKING EXPERIMENT — STANDALONE TEST")
    print("=" * 60)

    for label, path in [("QP", QP_PATH), ("Memo", MEMO_PATH), ("Submission", SUBMISSION)]:
        if not path.exists():
            print(f"[ERROR] {label} file not found: {path}")
            sys.exit(1)
        print(f"{label}: {path.name}")

    print("\n--- Reading files ---")
    qp_text   = read_docx_rich(QP_PATH)
    memo_text = read_docx_rich(MEMO_PATH)
    print(f"QP text length   : {len(qp_text)} chars")
    print(f"Memo text length : {len(memo_text)} chars")

    print("\n--- Parsing marksheet (heuristic, no LLM) ---")
    # Force heuristic by using a dummy backend that will fail fast and fall back
    parser = MarksheetParser(backend="ollama", model="none")

    try:
        result = parser.parse(qp_text, memo_text, program="word")
    except RuntimeError as exc:
        print(f"[ERROR] Parser failed: {exc}")
        sys.exit(1)

    print(f"Questions parsed : {len(result.questions)}")
    if result.warnings:
        print("Warnings:")
        for w in result.warnings:
            print(f"  - {w}".encode('cp1252', errors='replace').decode('cp1252'))

    print("\n--- Generated questions ---")
    for q in result.questions:
        print(f"  [{q['question_number']}] {q['description'][:70]}")
        print(f"        domain={q['domain']}  type={q['type']}  marks={q['marks']}")
        print(f"        target={q['target']}")
        print(f"        expected={str(q['expected']).encode('ascii', errors='replace').decode('ascii')}")

    task_definition = {
        "task_name": "CAT Test 1 Auto-Marked",
        "program": "word",
        "file": SUBMISSION.name,
        "total_marks": sum(int(q.get("marks", 1)) for q in result.questions),
        "questions": result.questions,
    }

    # Save the generated task JSON next to this script for inspection
    task_json_path = Path(__file__).parent / "generated_task.json"
    with task_json_path.open("w", encoding="utf-8") as f:
        json.dump(task_definition, f, indent=2, ensure_ascii=False)
    print(f"\nTask JSON saved to: {task_json_path}")

    print("\n--- Running marking engine ---")
    engine = MarkingEngine()
    session = engine.run_task(task_definition, SUBMISSION)

    print(f"\nSCORE: {session.score} / {session.total_marks}")
    print("\n--- Results ---")
    for r in session.results:
        status = "PASS" if r.passed else "FAIL"
        print(f"  [{status}] Q{r.question_number} ({r.marks}mk) — {r.description[:60]}")
        print(f"         Feedback: {r.feedback}")
        if r.details:
            for k, v in r.details.items():
                if k != "type":
                    print(f"         {k}: {v}")

    print("\n" + "=" * 60)
    print(f"FINAL SCORE: {session.score} / {session.total_marks}")
    print("=" * 60)


if __name__ == "__main__":
    main()
