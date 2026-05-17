"""Shared utilities for Word document checking."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
from zipfile import ZipFile

from ..checker_types import CheckerResult
from ..marking_experiment import resolve_theme_color_name

NAMESPACES = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

ALIGNMENT_MAP = {
    WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
    WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
    WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
    WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
    WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: "justify",
    WD_PARAGRAPH_ALIGNMENT.THAI_JUSTIFY: "justify",
}


def xml_find(element, xpath: str):
    return element.find(xpath, namespaces=NAMESPACES)


def normalize_hex_color(color: Optional[str]) -> Optional[str]:
    if not color:
        return None
    value = color.strip().upper().lstrip("#")
    return value if len(value) == 6 else value


def _xml_tree_from_oxml(oxml_element) -> Optional[etree._Element]:
    if oxml_element is None:
        return None
    return etree.fromstring(oxml_element.xml.encode("utf-8"))


def _read_docx_part(file_path: Path, part_name: str) -> Optional[str]:
    try:
        with ZipFile(file_path, "r") as zf:
            return zf.read(part_name).decode("utf-8")
    except Exception:
        return None


def _resolve_color_info(run, file_path: Path) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]:
    xml = run._r.xml
    tree = etree.fromstring(xml.encode("utf-8"))
    color_element = tree.find(".//w:color", namespaces=NAMESPACES)
    if color_element is None:
        return None, None, None, None
    theme_color = color_element.get(qn("w:themeColor"))
    theme_tint = color_element.get(qn("w:themeTint"))
    theme_shade = color_element.get(qn("w:themeShade"))
    value = color_element.get(qn("w:val"))
    if value:
        value = normalize_hex_color(value)
    return theme_color, theme_tint, theme_shade, value


def _match_color(expected: Any, actual_hex: Optional[str], actual_theme_name: Optional[str]) -> bool:
    expected_value = normalize_hex_color(str(expected)) if isinstance(expected, str) else None
    if expected_value and actual_hex:
        if expected_value == actual_hex:
            return True
    if isinstance(expected, str) and actual_theme_name:
        return expected.lower() == actual_theme_name.lower()
    return False


def _run_xml_tree(run) -> Optional[etree._Element]:
    if run is None or not hasattr(run, "_r"):
        return None
    try:
        return etree.fromstring(run._r.xml.encode("utf-8"))
    except Exception:
        return None


def _font_xml_element(run, tag_name: str):
    tree = _run_xml_tree(run)
    if tree is None:
        return None
    return tree.find(f".//w:{tag_name}", namespaces=NAMESPACES)


def _font_xml_bool(run, tag_name: str) -> bool:
    return _font_xml_element(run, tag_name) is not None


def _font_xml_val(run, tag_name: str) -> Optional[str]:
    element = _font_xml_element(run, tag_name)
    if element is None:
        return None
    return element.get(qn("w:val")) or None


def _resolve_font_theme(run) -> Optional[str]:
    tree = _run_xml_tree(run)
    if tree is None:
        return None
    fonts = tree.find(".//w:rFonts", namespaces=NAMESPACES)
    if fonts is None:
        return None
    for attr in ("asciiTheme", "hAnsiTheme", "csTheme", "eastAsiaTheme"):
        theme_val = fonts.get(qn(f"w:{attr}"))
        if theme_val:
            return theme_val
    return None


def _target_locator(target: Dict[str, Any]) -> Tuple[str, Any]:
    locator = target.get("locator")
    if isinstance(locator, dict):
        pair = next(iter(locator.items()))
        return pair[0], pair[1]
    return locator, target.get("value")


def _find_paragraphs(document: Document, target: Dict[str, Any]) -> List[Any]:
    locator, value = _target_locator(target)
    value = str(value) if value is not None else ""
    if locator == "paragraph_index":
        try:
            idx = int(value)
            # If the indexed paragraph has no formatting, scan all body paragraphs
            p = document.paragraphs[idx]
            return [p]
        except (IndexError, ValueError):
            return []
    if locator == "after_heading":
        for idx, paragraph in enumerate(document.paragraphs):
            if value.lower() in paragraph.text.lower():
                return document.paragraphs[idx + 1 : idx + 2]
        return []
    if locator == "style_name":
        return [p for p in document.paragraphs if p.style and p.style.name == value]
    if locator == "contains_text":
        return [p for p in document.paragraphs if value.lower() in p.text.lower()]
    if locator == "header_contains_text":
        section = document.sections[0]
        header = section.header
        if not header:
            return []
        # Match on plain text; if value is empty return all header paragraphs
        if value:
            return [p for p in header.paragraphs if value.lower() in p.text.lower()]
        return list(header.paragraphs)
    if locator == "footer_contains_text":
        section = document.sections[0]
        footer = section.footer
        if not footer:
            return []
        # Match on plain text first
        matched = [p for p in footer.paragraphs if value.lower() in p.text.lower()]
        if matched:
            return matched
        # Fall back: return all footer paragraphs (covers field-code-only footers)
        return list(footer.paragraphs)
    if locator == "starts_with":
        return [p for p in document.paragraphs if p.text.strip().lower().startswith(value.lower())]
    if locator == "document":
        return list(document.paragraphs)
    if locator == "near_text":
        return [p for p in document.paragraphs if value.lower() in p.text.lower()]
    return []


def _find_table(document: Document, target: Dict[str, Any]) -> Optional[Any]:
    locator, value = _target_locator(target)
    if locator == "table_index":
        try:
            return document.tables[int(value)]
        except (IndexError, ValueError):
            return None
    if locator == "table_near_text" or locator == "near_text":
        needle = str(value).lower()
        for paragraph in document.paragraphs:
            if needle in paragraph.text.lower():
                next_table = _find_table_after_paragraph(document, paragraph)
                if next_table is not None:
                    return next_table
        return None
    return None


def _find_table_after_paragraph(document: Document, paragraph) -> Optional[Any]:
    paragraphs = list(document.element.body.iterchildren())
    target_xml = paragraph._p
    found = False
    for element in paragraphs:
        if element is target_xml:
            found = True
            continue
        if found and element.tag == qn("w:tbl"):
            tbl = element
            from docx.table import Table
            return Table(tbl, document)
    return None


def _get_table_cell(table, row: int, col: int) -> Optional[Any]:
    try:
        return table.rows[row].cells[col]
    except (IndexError, ValueError):
        return None


def _cell_grid_span(cell) -> int:
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return 1
    grid_span = tcPr.find(qn("w:gridSpan"))
    if grid_span is not None and grid_span.get(qn("w:val")):
        return int(grid_span.get(qn("w:val")))
    return 1


def _cell_vmerge(cell) -> bool:
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return False
    vmerge = tcPr.find(qn("w:vMerge"))
    return vmerge is not None


def _find_run(paragraphs, text: Optional[str] = None):
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if not text or text.lower() in run.text.lower():
                return run
    return None


def _find_tables(document: Document, target: Dict[str, Any]) -> List[Any]:
    locator, value = _target_locator(target)
    if locator == "table_index":
        try:
            idx = int(value)
            return [document.tables[idx]]
        except (IndexError, ValueError):
            return []
    if locator == "table_near_text" or locator == "near_text":
        needle = str(value).lower()
        for paragraph in document.paragraphs:
            if needle in paragraph.text.lower():
                next_table = _find_table_after_paragraph(document, paragraph)
                if next_table is not None:
                    return [next_table]
        return []
    if locator == "all":
        return list(document.tables)
    return []


def _paragraph_num_pr(paragraph):
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    return xml_find(pPr, "w:numPr")


def _get_list_properties(paragraph, document):
    numPr = _paragraph_num_pr(paragraph)
    if numPr is None:
        return None

    numId_el = xml_find(numPr, "w:numId")
    ilvl_el = xml_find(numPr, "w:ilvl")
    if numId_el is None:
        return None

    num_id = numId_el.get(qn("w:val"))
    level = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0

    if not hasattr(document.part, "numbering_part") or document.part.numbering_part is None:
        return {"num_id": num_id, "level": level}
    numbering_xml = document.part.numbering_part.element.xml
    numbering_tree = etree.fromstring(numbering_xml.encode("utf-8"))
    num_nodes = numbering_tree.xpath(f".//w:num[@w:numId='{num_id}']", namespaces=NAMESPACES)
    if not num_nodes:
        return {"num_id": num_id, "level": level}

    abstract_id = num_nodes[0].xpath("./w:abstractNumId", namespaces=NAMESPACES)
    if not abstract_id:
        return {"num_id": num_id, "level": level}

    abstract_id = abstract_id[0].get(qn("w:val"))
    abstract_nodes = numbering_tree.xpath(f".//w:abstractNum[@w:abstractNumId='{abstract_id}']", namespaces=NAMESPACES)
    if not abstract_nodes:
        return {"num_id": num_id, "level": level}

    lvl_nodes = abstract_nodes[0].xpath(f".//w:lvl[@w:ilvl='{level}']", namespaces=NAMESPACES)
    if not lvl_nodes:
        return {"num_id": num_id, "level": level}

    num_fmt_el = lvl_nodes[0].find("./w:numFmt", namespaces=NAMESPACES)
    lvl_text_el = lvl_nodes[0].find("./w:lvlText", namespaces=NAMESPACES)
    num_fmt = num_fmt_el.get(qn("w:val")) if num_fmt_el is not None else None
    lvl_text = lvl_text_el.get(qn("w:val")) if lvl_text_el is not None else None

    if num_fmt is None:
        list_type = "unknown"
    elif num_fmt.lower() == "bullet":
        list_type = "bullet"
    elif num_fmt.lower() in {"decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman"}:
        list_type = "number"
    else:
        list_type = "multilevel"

    return {
        "num_id": num_id,
        "level": level,
        "num_fmt": num_fmt,
        "lvl_text": lvl_text,
        "type": list_type,
    }


def _match_run_font(run, expected: Dict[str, Any]) -> Tuple[bool, Dict[str, Any]]:
    actual: Dict[str, Any] = {}
    passed = True

    if "bold" in expected:
        actual["bold"] = bool(run.font.bold)
        passed = passed and actual["bold"] == bool(expected["bold"])
    if "italic" in expected:
        actual["italic"] = bool(run.font.italic)
        passed = passed and actual["italic"] == bool(expected["italic"])
    if "size" in expected:
        actual["size"] = run.font.size.pt if run.font.size else None
        passed = passed and float(actual["size"] or 0) == float(expected["size"])
    if "color" in expected:
        theme_color, theme_tint, theme_shade, value = _resolve_color_info(run, None)
        actual["color"] = value
        theme_name = resolve_theme_color_name(theme_color, theme_tint, theme_shade) if theme_color else None
        passed = passed and _match_color(expected["color"], value, theme_name)
    return passed, actual
