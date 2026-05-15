"""List rule checker for Word documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document

from ..checker_types import CheckerResult
from .word_utils import _find_paragraphs


def check_list_rule(rule: Dict[str, Any], file_path: Path) -> CheckerResult:
    document = Document(file_path)
    check_type = str(rule.get("type", ""))
    target = rule.get("target", {}) or {}
    expected = rule.get("expected")

    paragraphs = _find_paragraphs(document, target)
    if not paragraphs:
        return CheckerResult(passed=False, details={"reason": "List target not found."})
    paragraph = paragraphs[0]
    actual = None
    passed = False

    if check_type == "exists":
        # Check if paragraph has list formatting
        passed = paragraph.style.name.lower().startswith("list") or hasattr(paragraph, "_p") and paragraph._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr") is not None
        actual = "List found" if passed else "No list"

    elif check_type == "type":
        # Check list type (numbered, bulleted)
        list_type = _get_list_type(paragraph)
        actual = list_type
        passed = list_type == str(expected).lower()

    elif check_type == "level":
        # Check indentation level
        level = _get_list_level(paragraph)
        actual = level
        try:
            passed = level == int(expected)
        except Exception:
            passed = False

    elif check_type == "style":
        # Check list style name
        style_name = paragraph.style.name if paragraph.style else ""
        actual = style_name
        passed = str(expected).lower() in style_name.lower()

    elif check_type == "count":
        # Count list items in document
        list_items = _count_list_items(document)
        actual = list_items
        try:
            passed = list_items == int(expected)
        except Exception:
            passed = False

    elif check_type == "continuation":
        # Check if list continues properly
        passed, actual = _check_list_continuation(document, expected)

    else:
        return CheckerResult(passed=False, details={"reason": "Unsupported list check."})

    return CheckerResult(passed=passed, actual=actual, details={"type": check_type})


def _get_list_type(paragraph) -> str:
    """Determine if paragraph is numbered or bulleted list."""
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        p_xml = etree.fromstring(paragraph._p.xml.encode("utf-8"))
        num_pr = p_xml.find(f".//{{{WNS}}}numPr")
        if num_pr is not None:
            num_id = num_pr.find(f"{{{WNS}}}numId")
            if num_id is not None:
                return "numbered"
        # Check for bullet style
        if paragraph.style and "bullet" in paragraph.style.name.lower():
            return "bulleted"
        # Check for common list styles
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if "list" in style_name:
            if "number" in style_name or "decimal" in style_name:
                return "numbered"
            else:
                return "bulleted"
    except Exception:
        pass
    return "none"


def _get_list_level(paragraph) -> int:
    """Get the indentation level of a list item."""
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        p_xml = etree.fromstring(paragraph._p.xml.encode("utf-8"))
        num_pr = p_xml.find(f".//{{{WNS}}}numPr")
        if num_pr is not None:
            ilvl = num_pr.find(f"{{{WNS}}}ilvl")
            if ilvl is not None:
                val = ilvl.get(f"{{{WNS}}}val", "0")
                return int(val)
        # Fallback to indentation
        indent = paragraph.paragraph_format.left_indent
        if indent:
            # Rough estimate: 0.5 inch per level
            level = int(indent.pt / 36)
            return level
    except Exception:
        pass
    return 0


def _count_list_items(document) -> int:
    """Count total list items in document."""
    count = 0
    for para in document.paragraphs:
        if _get_list_type(para) != "none":
            count += 1
    return count


def _check_list_continuation(document, expected: Any) -> tuple[bool, Any]:
    """Check if lists continue properly without breaks."""
    list_sequences = []
    current_list = []
    current_type = None

    for para in document.paragraphs:
        list_type = _get_list_type(para)
        if list_type != "none":
            if list_type == current_type:
                current_list.append(para)
            else:
                if current_list:
                    list_sequences.append((current_type, current_list))
                current_list = [para]
                current_type = list_type
        else:
            if current_list:
                list_sequences.append((current_type, current_list))
                current_list = []
                current_type = None

    if current_list:
        list_sequences.append((current_type, current_list))

    # Check for continuity
    broken_lists = 0
    for list_type, items in list_sequences:
        if len(items) < 2:
            broken_lists += 1

    actual = {"sequences": len(list_sequences), "broken": broken_lists}
    if isinstance(expected, dict):
        min_sequences = expected.get("min_sequences", 1)
        max_broken = expected.get("max_broken", 0)
        passed = len(list_sequences) >= min_sequences and broken_lists <= max_broken
    else:
        passed = broken_lists == 0

    return passed, actual