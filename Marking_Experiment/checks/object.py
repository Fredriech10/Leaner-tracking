"""Object rule checker for Word documents (images, shapes, etc.)."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from docx import Document

from ..checker_types import CheckerResult


def check_object_rule(rule: Dict[str, Any], file_path: Path) -> CheckerResult:
    document = Document(file_path)
    check_type = str(rule.get("type", ""))
    target = rule.get("target", {}) or {}
    expected = rule.get("expected")

    actual = None
    passed = False

    if check_type == "image_exists":
        images = _find_images(document)
        actual = len(images)
        try:
            min_count = expected.get("min", 1) if isinstance(expected, dict) else int(expected)
            passed = actual >= min_count
        except Exception:
            passed = actual > 0

    elif check_type == "image_count":
        images = _find_images(document)
        actual = len(images)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "image_size":
        images = _find_images(document)
        if images:
            # Check first image
            img_data = images[0]
            actual = {"width": img_data.get("width"), "height": img_data.get("height")}
            if isinstance(expected, dict):
                exp_width = expected.get("width")
                exp_height = expected.get("height")
                tolerance = expected.get("tolerance", 10)  # pixels
                passed = (abs(actual["width"] - exp_width) <= tolerance and
                         abs(actual["height"] - exp_height) <= tolerance)
        else:
            passed = False
            actual = "No images found"

    elif check_type == "shape_exists":
        shapes = _find_shapes(document)
        actual = len(shapes)
        try:
            min_count = expected.get("min", 1) if isinstance(expected, dict) else int(expected)
            passed = actual >= min_count
        except Exception:
            passed = actual > 0

    elif check_type == "shape_type":
        shapes = _find_shapes(document)
        if shapes:
            shape_types = [s.get("type") for s in shapes]
            actual = shape_types
            if isinstance(expected, list):
                passed = all(st in shape_types for st in expected)
            else:
                passed = str(expected) in shape_types
        else:
            passed = False
            actual = "No shapes found"

    elif check_type == "chart_exists":
        charts = _find_charts(document)
        actual = len(charts)
        passed = actual > 0 if expected else actual == 0

    elif check_type == "table_of_contents":
        toc_found = _find_table_of_contents(document)
        actual = "Found" if toc_found else "Not found"
        passed = toc_found == bool(expected)

    elif check_type == "hyperlink_count":
        hyperlinks = _find_hyperlinks(document)
        actual = len(hyperlinks)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "footnote_count":
        footnotes = _find_footnotes(document)
        actual = len(footnotes)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    elif check_type == "endnote_count":
        endnotes = _find_endnotes(document)
        actual = len(endnotes)
        try:
            passed = actual == int(expected)
        except Exception:
            passed = False

    else:
        return CheckerResult(passed=False, details={"reason": "Unsupported object check."})

    return CheckerResult(passed=passed, actual=actual, details={"type": check_type})


def _find_images(document) -> list[Dict[str, Any]]:
    """Find all images in document."""
    images = []
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    VNS = "urn:schemas-microsoft-com:vml"
    try:
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                images.append({"id": rel.rId, "type": "inline_image"})
        # Also check for VML shapes with images
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            vml_images = p_xml.findall(f".//{{{VNS}}}imagedata")
            for img in vml_images:
                images.append({"id": img.get("id"), "type": "vml_image"})
    except Exception:
        pass
    return images


def _find_shapes(document) -> list[Dict[str, Any]]:
    """Find all shapes in document."""
    shapes = []
    from lxml import etree
    VNS = "urn:schemas-microsoft-com:vml"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            vml_shapes = p_xml.findall(f".//{{{VNS}}}shape")
            for shape in vml_shapes:
                shape_type = shape.get("type", "")
                shapes.append({"type": shape_type})
    except Exception:
        pass
    return shapes


def _find_charts(document) -> list[Dict[str, Any]]:
    """Find all charts in document."""
    charts = []
    from lxml import etree
    CNS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
    try:
        for rel in document.part.rels.values():
            if "chart" in rel.reltype:
                charts.append({"id": rel.rId})
    except Exception:
        pass
    return charts


def _find_table_of_contents(document) -> bool:
    """Check if document has a table of contents."""
    toc_found = False
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            # Look for TOC field codes
            fld_simple = p_xml.find(f".//{{{WNS}}}fldSimple")
            if fld_simple is not None:
                instr = fld_simple.get(f"{{{WNS}}}instr", "")
                if "TOC" in instr.upper():
                    toc_found = True
                    break
            # Look for complex fields
            fld_char = p_xml.find(f".//{{{WNS}}}fldChar")
            if fld_char is not None:
                # This is more complex to parse fully
                toc_found = True
                break
    except Exception:
        pass
    return toc_found


def _find_hyperlinks(document) -> list[Dict[str, Any]]:
    """Find all hyperlinks in document."""
    hyperlinks = []
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            hyperlinks_el = p_xml.findall(f".//{{{WNS}}}hyperlink")
            for link in hyperlinks_el:
                hyperlinks.append({"id": link.get("{{{WNS}}}id", "")})
    except Exception:
        pass
    return hyperlinks


def _find_footnotes(document) -> list[Dict[str, Any]]:
    """Find all footnotes in document."""
    footnotes = []
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            footnote_refs = p_xml.findall(f".//{{{WNS}}}footnoteReference")
            for ref in footnote_refs:
                footnotes.append({"id": ref.get("{{{WNS}}}id", "")})
    except Exception:
        pass
    return footnotes


def _find_endnotes(document) -> list[Dict[str, Any]]:
    """Find all endnotes in document."""
    endnotes = []
    from lxml import etree
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        for para in document.paragraphs:
            p_xml = etree.fromstring(para._p.xml.encode("utf-8"))
            endnote_refs = p_xml.findall(f".//{{{WNS}}}endnoteReference")
            for ref in endnote_refs:
                endnotes.append({"id": ref.get("{{{WNS}}}id", "")})
    except Exception:
        pass
    return endnotes