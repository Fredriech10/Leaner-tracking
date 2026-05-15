"""LLM-based marksheet parser for Marking Experiment."""

from __future__ import annotations

import json
import os
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from .marking_experiment import MARKING_SCHEMA


@dataclass
class ParseResult:
    questions: List[Dict[str, Any]]
    warnings: Optional[List[str]] = None

    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


class MarksheetParser:
    def __init__(self, model: Optional[str] = None, backend: str = "openai") -> None:
        self.backend = backend.lower()
        if model is None:
            self.model = "gpt-4o-mini" if self.backend == "openai" else "phi3:instruct"
        else:
            self.model = model
        self.openai = None

    def _load_openai_module(self):
        try:
            import openai

            return openai
        except ImportError:
            return None

    def parse(self, question_paper_text: str, marksheet_text: str, program: str = "word", filename: Optional[str] = None) -> ParseResult:
        # If the caller provided a DOCX path, rebuild rich text with table markers for reliable heuristic fallback.
        if program.lower() == "word" and filename:
            try:
                file_path = Path(filename)
                if file_path.exists() and file_path.suffix.lower() == ".docx":
                    if not self._is_rich_docx_text(marksheet_text):
                        marksheet_text = self._read_docx_text_with_tables(file_path)
            except Exception:
                pass

        # Always try LLM first
        try:
            prompt = self._build_prompt(question_paper_text, marksheet_text, program, filename)
            if self.backend == "openai":
                response_text = self._call_openai(prompt)
            elif self.backend == "ollama":
                response_text = self._call_ollama(prompt)
            else:
                raise ValueError(f"Unsupported backend: {self.backend}. Use 'openai' or 'ollama'.")
            questions = self._extract_json(response_text)
            validation_warnings = self._validate_questions(questions)
            return ParseResult(questions=questions, warnings=validation_warnings)
        except Exception as llm_error:
            # Fall back to heuristic parser only if LLM fails
            questions = self._parse_memo_tables(marksheet_text, program)
            if questions:
                validation_warnings = self._validate_questions(questions)
                validation_warnings.append(f"LLM failed ({llm_error}); parsed heuristically from memo tables.")
                return ParseResult(questions=questions, warnings=validation_warnings)
            raise RuntimeError(
                f"LLM parsing failed and heuristic fallback found no questions. LLM error: {llm_error}"
            )

    def build_task_definition(
        self,
        question_paper_text: str,
        marksheet_text: str,
        program: str = "word",
        filename: Optional[str] = None,
        task_name: Optional[str] = None,
        submission_filename: Optional[str] = None,
    ) -> Dict[str, Any]:
        result = self.parse(question_paper_text, marksheet_text, program, filename)
        if task_name is None:
            task_name = "Generated Marking Task"

        questions: List[Dict[str, Any]] = []
        for q in result.questions:
            if isinstance(q.get("rules"), list):
                questions.append(q)
                continue
            questions.append(
                {
                    "question_number": q.get("question_number", ""),
                    "description": q.get("description", ""),
                    "target": q.get("target", {}),
                    "rules": [
                        {
                            "domain": q.get("domain", ""),
                            "type": q.get("type", ""),
                            "target": q.get("target", {}),
                            "expected": q.get("expected"),
                            "marks": int(q.get("marks", 1)),
                            "description": q.get("description", ""),
                        }
                    ],
                }
            )

        total_marks = sum(int(rule.get("marks", 1)) for question in questions for rule in question.get("rules", []))
        return {
            "task_name": task_name,
            "program": program,
            "file": submission_filename or "",
            "total_marks": total_marks,
            "questions": questions,
        }

    def _parse_memo_tables(self, marksheet_text: str, program: str) -> List[Dict[str, Any]]:
        if program.lower() != "word":
            return []

        table_blocks = self._extract_table_blocks(marksheet_text)
        for rows in table_blocks:
            if not rows:
                continue
            header = rows[0]
            if len(header) >= 2 and any("no." in cell.lower() for cell in header) and any("criteria" in cell.lower() for cell in header):
                parsed = self._parse_criteria_table(rows[1:])
                if parsed:
                    return parsed
        return []

    def _is_rich_docx_text(self, text: str) -> bool:
        return "TABLE " in text or "\t" in text or text.strip().startswith("[P")

    def _read_docx_text_with_tables(self, file_path: Path) -> str:
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError as exc:
            raise RuntimeError("python-docx is required for reading DOCX files. Install it with 'pip install python-docx'.") from exc

        align_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        }

        document = Document(str(file_path))
        parts: List[str] = []
        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            hints = []
            align = align_map.get(paragraph.alignment) if paragraph.alignment is not None else None
            if align:
                hints.append(f"align={align}")
            fmt = paragraph.paragraph_format
            if fmt.line_spacing is not None:
                hints.append(f"line_spacing={fmt.line_spacing}")
            if fmt.space_before is not None:
                try:
                    hints.append(f"space_before={fmt.space_before.pt}pt")
                except Exception:
                    pass
            if fmt.space_after is not None:
                try:
                    hints.append(f"space_after={fmt.space_after.pt}pt")
                except Exception:
                    pass
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.size:
                        try:
                            hints.append(f"font_size={run.font.size.pt}pt")
                        except Exception:
                            pass
                    if run.font.bold:
                        hints.append("bold")
                    if run.font.italic:
                        hints.append("italic")
                    if run.font.name:
                        hints.append(f"font={run.font.name}")
                    break
            style_name = paragraph.style.name if paragraph.style else ""
            hint_str = f" [{', '.join(hints)}]" if hints else ""
            style_str = f" (style: {style_name})" if style_name else ""
            parts.append(f"[P{idx}]{style_str} {text}{hint_str}")

        for table_idx, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("\t".join(cells))
            parts.append(f"TABLE {table_idx}:\n" + "\n".join(rows))

        return "\n\n".join(parts)

    def _extract_table_blocks(self, text: str) -> List[List[List[str]]]:
        blocks: List[List[List[str]]] = []
        current: List[List[str]] = []
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.upper().startswith("TABLE"):
                if current:
                    blocks.append(current)
                    current = []
                continue
            if "\t" in line:
                current.append([cell.strip() for cell in line.split("\t")])
            elif current and not stripped:
                blocks.append(current)
                current = []
        if current:
            blocks.append(current)
        return blocks

    def _parse_criteria_table(self, rows: List[List[str]]) -> List[Dict[str, Any]]:
        questions: List[Dict[str, Any]] = []
        section_number = None
        section_index = 1
        TOTAL_KEYWORDS = {"total", "totaal", "grand total"}
        # Section label keywords — rows that name a category, not a check
        SECTION_LABEL_KEYWORDS = {
            "page setup", "margins", "hyphenation", "alignment", "line spacing",
            "indentation", "borders and shading", "page border", "watermark",
            "special formatting", "drop cap", "header", "footer",
            "line and paragraph spacing",
        }

        pending: Optional[Dict[str, Any]] = None  # holds a partial check waiting for a measurement row

        for row in rows:
            if not any(cell.strip() for cell in row):
                continue
            first  = row[0].strip() if len(row) >= 1 else ""
            second = row[1].strip() if len(row) >= 2 else ""
            third  = row[2].strip() if len(row) >= 3 else ""

            # Section header: non-empty first col
            if first:
                # flush any pending partial check
                if pending:
                    questions.append(pending)
                    pending = None
                section_number = first
                section_index = 1
                continue

            # Check row: empty first col, description in second col
            if not first and second:
                if second.strip().lower() in TOTAL_KEYWORDS:
                    if pending:
                        questions.append(pending)
                        pending = None
                    continue

                # Skip pure section-label rows (no marks)
                if second.strip().lower() in SECTION_LABEL_KEYWORDS and not third.strip():
                    if pending:
                        questions.append(pending)
                        pending = None
                    continue

                marks = self._parse_mark_value(third)
                question_number = f"{section_number}.{section_index}" if section_number else str(len(questions) + 1)

                # Detect measurement-only follow-up rows (e.g. "Measurement set to 0.5 cm")
                # and merge them into the previous pending check
                if pending and self._is_measurement_row(second):
                    self._merge_measurement(pending, second)
                    pending["marks"] = pending.get("marks", 0) + marks
                    questions.append(pending)
                    pending = None
                    section_index += 1
                    continue

                # Flush any unmerged pending
                if pending:
                    questions.append(pending)
                    pending = None

                question = self._build_question_from_description(question_number, second, marks)
                if question:
                    # If this looks like it will have a follow-up measurement row, hold it
                    if self._expects_measurement_followup(second):
                        pending = question
                    else:
                        questions.append(question)
                        section_index += 1

        if pending:
            questions.append(pending)
        return questions

    def _is_measurement_row(self, text: str) -> bool:
        t = text.lower()
        return (
            t.startswith("measurement") or
            ("set to" in t and re.search(r"\d+", t) is not None and "spacing" not in t and "colour" not in t)
        )

    def _expects_measurement_followup(self, text: str) -> bool:
        t = text.lower()
        return (
            "spacing set to" in t or
            "first line indent applied" in t or
            ("indent" in t and "measurement" not in t and re.search(r"\d+", t) is None)
        )

    def _merge_measurement(self, question: Dict[str, Any], measurement_text: str) -> None:
        """Merge a measurement follow-up row into an existing question."""
        t = measurement_text.lower()
        val_cm = re.search(r"(\d+(?:\.\d+)?)\s*cm", t)
        val_pt = re.search(r"(\d+(?:\.\d+)?)\s*pt", t)
        if question.get("type") == "line_spacing" and isinstance(question.get("expected"), dict):
            if val_pt:
                question["expected"]["value"] = float(val_pt.group(1))
                question["expected"]["unit"] = "pt"
        elif question.get("type") == "first_line_indent":
            if val_cm:
                question["expected"] = float(val_cm.group(1))

    def _parse_mark_value(self, text: str) -> int:
        if not text:
            return 1
        match = re.search(r"(\d+)", text)
        return int(match.group(1)) if match else 1

    def _build_question_from_description(self, question_number: str, description: str, marks: int) -> Optional[Dict[str, Any]]:
        desc = description.strip()
        desc_lower = desc.lower()
        target: Dict[str, Any] = {"locator": "document"}
        expected: Any = True
        domain = "document"
        check_type = "header_text"

        # --- document-level checks ---
        if "paper size" in desc_lower or "document size" in desc_lower or "page size" in desc_lower:
            check_type = "paper_size"
            match = re.search(r"a4|letter|legal", desc_lower)
            expected = match.group(0).upper() if match else "A4"
        elif "orientation" in desc_lower:
            check_type = "orientation"
            expected = "portrait" if "portrait" in desc_lower else "landscape"
        elif "margin" in desc_lower:
            check_type = "margins"
            values = re.findall(r"(\d+(?:\.\d+)?)\s*cm", desc_lower)
            if values:
                expected = {
                    "top": float(values[0]),
                    "bottom": float(values[1]) if len(values) > 1 else float(values[0]),
                    "left": float(values[2]) if len(values) > 2 else float(values[0]),
                    "right": float(values[3]) if len(values) > 3 else float(values[0]),
                }
            else:
                expected = {"all": True}
        elif "first page only" in desc_lower:
            check_type = "page_border"
            expected = {"style": "double", "first_page_only": True}
        elif "page border" in desc_lower or "double-line" in desc_lower or "double line" in desc_lower or "border applied" in desc_lower or ("style" in desc_lower and "double" in desc_lower):
            check_type = "page_border"
            style_match = re.search(r"double|single|thick|dashed|dotted", desc_lower)
            color_match = re.search(r"colour[:\s]+(blue|red|green|black|yellow)|color[:\s]+(blue|red|green|black|yellow)", desc_lower)
            expected = {"style": style_match.group(0) if style_match else "double"}
            if color_match:
                expected["color"] = (color_match.group(1) or color_match.group(2))
        elif "watermark" in desc_lower or "cat test" in desc_lower:
            check_type = "watermark"
            text_match = re.search(r'"([^"]+)"', desc)
            if text_match:
                expected = text_match.group(1)
            elif "cat test" in desc_lower:
                expected = "CAT TEST"
            elif "diagonal" in desc_lower:
                expected = "Diagonal"
            else:
                expected = True
        elif "title:" in desc_lower or ("title" in desc_lower and "entered" in desc_lower and "style" not in desc_lower):
            domain = "paragraph_formatting"
            check_type = "contains_text"
            # Extract text after "Title:" up to next word boundary or end
            colon_match = re.search(r'title:\s*([A-Z0-9][A-Z0-9 ]+?)(?:\s+correctly|\s+entered|$)', desc, re.IGNORECASE)
            text_match = re.search(r'"([^"]+)"', desc)
            word = colon_match.group(1).strip() if colon_match else (text_match.group(1) if text_match else "")
            target = {"locator": "document"}
            expected = {"text": word}
        elif "subtitle:" in desc_lower or ("subtitle" in desc_lower and "entered" in desc_lower):
            domain = "paragraph_formatting"
            check_type = "contains_text"
            colon_match = re.search(r'subtitle:\s*([A-Z0-9][A-Z0-9 ]+?)(?:\s+correctly|\s+entered|$)', desc, re.IGNORECASE)
            text_match = re.search(r'"([^"]+)"', desc)
            word = colon_match.group(1).strip() if colon_match else (text_match.group(1) if text_match else "")
            target = {"locator": "document"}
            expected = {"text": word}
        elif "date" in desc_lower and "content control" in desc_lower:
            domain = "document"
            check_type = "contains_date"
            target = {"locator": "document"}
            expected = {"format": "date"}
        elif "header" in desc_lower and "name" in desc_lower and "right" in desc_lower:
            check_type = "header_alignment"
            expected = "right"
        elif "header" in desc_lower and "name" in desc_lower:
            check_type = "header_text"
            expected = True
        elif "header" in desc_lower and ("right" in desc_lower or "align" in desc_lower):
            check_type = "header_alignment"
            expected = "right" if "right" in desc_lower else "center" if "center" in desc_lower or "centre" in desc_lower else "left"
        elif "name" in desc_lower and "surname" in desc_lower and "right" in desc_lower:
            check_type = "header_alignment"
            expected = "right"
        elif "name" in desc_lower and "surname" in desc_lower:
            check_type = "header_text"
            expected = True
        elif "footer" in desc_lower and ("page number" in desc_lower or "automatic" in desc_lower):
            check_type = "footer_text"
            expected = True
        elif "page x of y" in desc_lower:
            check_type = "footer_text"
            expected = "Page X of Y"
        elif "footer" in desc_lower and ("center" in desc_lower or "centre" in desc_lower):
            check_type = "footer_alignment"
            expected = "center"
        elif "hyphenation" in desc_lower:
            check_type = "hyphenation"
            expected = True
        elif "colour set to" in desc_lower or "color set to" in desc_lower or ("colour" in desc_lower and "watermark" not in desc_lower and "border" not in desc_lower):
            check_type = "watermark"
            color_match = re.search(r"(blue|red|green|yellow|black|white|gray|grey)", desc_lower)
            expected = color_match.group(1).capitalize() if color_match else True
        elif "layout set to diagonal" in desc_lower or ("diagonal" in desc_lower and "layout" in desc_lower):
            check_type = "watermark"
            expected = "Diagonal"

        # --- paragraph_formatting checks ---
        elif any(w in desc_lower for w in ("align", "justified", "centred", "centered", "left align", "right align")):
            domain = "paragraph_formatting"
            check_type = "alignment"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "paragraph_index", "value": 2}
            if "justify" in desc_lower or "justified" in desc_lower:
                expected = "justify"
            elif "center" in desc_lower or "centred" in desc_lower:
                expected = "center"
            elif "right" in desc_lower:
                expected = "right"
            else:
                expected = "left"
        elif "line spacing" in desc_lower or "line space" in desc_lower or ("spacing" in desc_lower and "exactly" in desc_lower) or ("spacing set to" in desc_lower):
            domain = "paragraph_formatting"
            check_type = "line_spacing"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "paragraph_index", "value": 2}
            rule = "exact"
            if "exactly" in desc_lower or "exact" in desc_lower:
                rule = "exact"
            elif "at least" in desc_lower:
                rule = "atLeast"
            elif "multiple" in desc_lower or "double" in desc_lower:
                rule = "multiple"
            val_match = re.search(r"(\d+(?:\.\d+)?)\s*pt", desc_lower)
            val_match2 = re.search(r"(\d+(?:\.\d+)?)\s*lines?", desc_lower)
            if val_match:
                expected = {"rule": rule, "value": float(val_match.group(1)), "unit": "pt"}
            elif val_match2:
                expected = {"rule": rule, "value": float(val_match2.group(1)), "unit": "lines"}
            else:
                expected = {"rule": rule, "value": 0.0, "unit": "pt"}
        elif "spacing" in desc_lower and "before" in desc_lower:
            domain = "paragraph_formatting"
            check_type = "space_before"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            val_match = re.search(r"(\d+(?:\.\d+)?)\s*pt", desc_lower)
            expected = float(val_match.group(1)) if val_match else 12.0
        elif "spacing" in desc_lower and "after" in desc_lower:
            domain = "paragraph_formatting"
            check_type = "space_after"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            val_match = re.search(r"(\d+(?:\.\d+)?)\s*pt", desc_lower)
            expected = float(val_match.group(1)) if val_match else 0.0
        elif "first line indent" in desc_lower or "first line" in desc_lower and "indent" in desc_lower:
            domain = "paragraph_formatting"
            check_type = "first_line_indent"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under|starting with)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            val_match = re.search(r"(\d+(?:\.\d+)?)\s*cm", desc_lower)
            expected = float(val_match.group(1)) if val_match else 0.5
        elif "indent" in desc_lower:
            domain = "paragraph_formatting"
            check_type = "first_line_indent"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under|starting with)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            val_match = re.search(r"(\d+(?:\.\d+)?)\s*cm", desc_lower)
            expected = float(val_match.group(1)) if val_match else 0.5
        elif "border" in desc_lower and "paragraph" not in desc_lower:
            domain = "paragraph_formatting"
            check_type = "border"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            width_match = re.search(r"(\d+(?:\.\d+)?)\s*pt", desc_lower)
            color_match = re.search(r"(blue|red|green|black|yellow|white|gray|grey)", desc_lower)
            style_match = re.search(r"(double|single|thick|dashed|dotted|solid)", desc_lower)
            expected = {
                "sides": "all",
                "style": style_match.group(1) if style_match else "single",
                "color": color_match.group(1) if color_match else None,
                "width_pt": float(width_match.group(1)) if width_match else 1.0,
            }
        elif "shading" in desc_lower or "fill" in desc_lower:
            domain = "paragraph_formatting"
            check_type = "shading"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under|same paragraph|same)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            color_match = re.search(r"(light grey|light gray|grey|gray|blue|red|green|yellow|black|white)", desc_lower)
            expected = {"color": color_match.group(1) if color_match else "any"}
        elif "drop cap" in desc_lower or "dropcap" in desc_lower:
            domain = "paragraph_formatting"
            check_type = "drop_cap"
            heading_match = re.search(r'(?:under|after|below|located under|paragraph under|first letter under)\s+(?:the\s+)?(?:heading\s+)?["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "document"}
            lines_match = re.search(r"(\d+)\s*lines?", desc_lower)
            expected = {"lines": int(lines_match.group(1)) if lines_match else 3}

        # --- font checks ---
        elif any(w in desc_lower for w in ("font size", "font colour", "font color", "bold", "italic")):
            domain = "font"
            heading_match = re.search(r'(?:under|after|below|in|of)\s+["\']?([A-Za-z0-9 ]+)["\']?', desc)
            target = {"locator": "after_heading", "value": heading_match.group(1).strip()} if heading_match else {"locator": "paragraph_index", "value": 0}
            if "font size" in desc_lower:
                check_type = "size"
                val_match = re.search(r"(\d+(?:\.\d+)?)\s*pt", desc_lower)
                expected = float(val_match.group(1)) if val_match else 12.0
            elif "font colour" in desc_lower or "font color" in desc_lower:
                check_type = "color"
                color_match = re.search(r"#?([0-9A-Fa-f]{6})", desc)
                expected = color_match.group(1).upper() if color_match else "000000"
            elif "bold" in desc_lower:
                check_type = "bold"
                expected = "not bold" not in desc_lower
            elif "italic" in desc_lower:
                check_type = "italic"
                expected = "not italic" not in desc_lower

        elif "text \"adventure\" is bold and red" in desc_lower:
            domain = "font"
            check_type = "bold_and_color"
            target = {"locator": "contains_text", "value": "Adventure"}
            expected = {"bold": True, "color": "red"}

        # --- style checks ---
        elif "style" in desc_lower and any(w in desc_lower for w in ("applied", "apply", "title", "heading")):
            domain = "advanced"
            check_type = "style_applied"
            style_match = re.search(r'(?:title|heading\s*\d?|normal|body text)', desc_lower)
            text_match = re.search(r'"([^"]+)"', desc)
            target = {"locator": "contains_text", "value": text_match.group(1)} if text_match else {"locator": "document"}
            expected = {"style": style_match.group(0).strip().title() if style_match else "Title"}

        # --- find and replace checks ---
        elif "replaced" in desc_lower or ("replace" in desc_lower and "find" in desc_lower):
            domain = "paragraph_formatting"
            check_type = "contains_text"
            # Check for the replacement word (second quoted string, or after "with")
            with_match = re.search(r'with\s+"([^"]+)"', desc, re.IGNORECASE)
            text_match = re.findall(r'"([^"]+)"', desc)
            word = with_match.group(1) if with_match else (text_match[-1] if text_match else "")
            target = {"locator": "document"}
            expected = {"text": word}

        # --- image/object checks ---
        elif "width" in desc_lower and "cm" in desc_lower:
            domain = "object"
            check_type = "image_width"
            val_match = re.search(r"(\d+(?:\.\d+)?)\s*cm", desc_lower)
            target = {"locator": "document"}
            expected = {"width_cm": float(val_match.group(1)) if val_match else 5.0}

        elif "border" in desc_lower and any(w in desc_lower for w in ("image", "picture", "photo", "flag", "pt", "applied")):
            domain = "object"
            check_type = "image_border"
            width_match = re.search(r"(\d+(?:\.\d+)?)\s*pt", desc_lower)
            color_match = re.search(r"(blue|red|green|black|yellow|white)", desc_lower)
            target = {"locator": "document"}
            expected = {
                "width_pt": float(width_match.group(1)) if width_match else 1.0,
                "color": color_match.group(1) if color_match else None,
            }

        elif "cropped" in desc_lower or "crop" in desc_lower:
            domain = "object"
            check_type = "image_crop"
            shape_match = re.search(r"(oval|circle|rectangle|rounded|diamond|triangle)", desc_lower)
            target = {"locator": "document"}
            expected = {"shape": shape_match.group(1) if shape_match else "oval"}

        # --- SmartArt checks ---
        elif "smartart" in desc_lower or "smart art" in desc_lower:
            domain = "object"
            target = {"locator": "document"}
            if "text" in desc_lower or "stage" in desc_lower or "node" in desc_lower or "entered" in desc_lower:
                check_type = "smartart_text"
                text_match = re.search(r'"([^"]+)"', desc)
                expected = {"contains": text_match.group(1) if text_match else ""}
            elif "color" in desc_lower or "colour" in desc_lower or "colorful" in desc_lower or "colourful" in desc_lower:
                check_type = "smartart_color"
                expected = {"scheme": "colorful"}
            else:
                check_type = "smartart"
                type_match = re.search(r"(basic process|hierarchy|cycle|list|process|relationship|matrix|pyramid)", desc_lower)
                expected = {"type": type_match.group(1) if type_match else "basic process"}

        elif "colorful" in desc_lower or "colourful" in desc_lower:
            domain = "object"
            check_type = "smartart_color"
            target = {"locator": "document"}
            expected = {"scheme": "colorful"}

        elif "correct text" in desc_lower and any(w in desc_lower for w in ("plan", "explore", "reflect", "stage", "entered", "order")):
            domain = "object"
            check_type = "smartart_text"
            target = {"locator": "document"}
            # Extract text inside parentheses as the expected content
            paren_match = re.search(r'\(([^)]+)\)', desc)
            expected = {"contains": paren_match.group(1) if paren_match else ""}

        # --- bullet/list checks ---
        elif "wingdings" in desc_lower or "custom bullet" in desc_lower or ("bullet" in desc_lower and "character" in desc_lower):
            domain = "list"
            check_type = "bullet_char"
            target = {"locator": "style_name", "value": "List Paragraph"}
            code_match = re.search(r"code\s*(\d+)", desc_lower)
            if code_match:
                char_code = int(code_match.group(1))
                expected = chr(0xF000 + char_code)
            else:
                expected = "wingdings"

        elif "correct symbol" in desc_lower or ("code" in desc_lower and "symbol" in desc_lower) or ("code" in desc_lower and "checkmark" in desc_lower):
            domain = "list"
            check_type = "bullet_char"
            target = {"locator": "style_name", "value": "List Paragraph"}
            code_match = re.search(r"code\s*(\d+)", desc_lower)
            if code_match:
                char_code = int(code_match.group(1))
                expected = chr(0xF000 + char_code)
            else:
                expected = "wingdings"

        # --- table checks ---
        elif "merge" in desc_lower or "merged" in desc_lower:
            domain = "table"
            check_type = "merge_horizontal"
            target = {"locator": "table_index", "value": 0}
            row_match = re.search(r"row\s*(\d+)", desc_lower)
            col_match = re.findall(r"col(?:umn)?\s*(\d+)", desc_lower)
            row = int(row_match.group(1)) - 1 if row_match else 0
            col_start = int(col_match[0]) - 1 if len(col_match) > 0 else 0
            col_end = int(col_match[1]) - 1 if len(col_match) > 1 else col_start + 1
            expected = {"row": row, "col_start": col_start, "col_end": col_end}
        elif "cell" in desc_lower and "text" in desc_lower:
            domain = "table"
            check_type = "cell_text"
            target = {"locator": "table_index", "value": 0}
            text_match = re.search(r'"([^"]+)"', desc)
            expected = {"row": 0, "col": 0, "text": text_match.group(1) if text_match else "", "tolerance": True}

        # --- list checks ---
        elif "bullet" in desc_lower or "numbered list" in desc_lower or "list" in desc_lower:
            domain = "list"
            check_type = "list_style"
            target = {"locator": "document"}
            expected = {"type": "bullet", "level": 0} if "bullet" in desc_lower else {"type": "number", "level": 0}

        # --- advanced checks ---
        elif "bookmark" in desc_lower:
            domain = "advanced"
            check_type = "bookmark"
            target = {"locator": "document"}
            name_match = re.search(r'"([^"]+)"', desc)
            expected = {"name": name_match.group(1) if name_match else "bookmark1"}
        elif "bibliography" in desc_lower or "references" in desc_lower:
            domain = "advanced"
            check_type = "bibliography"
            target = {"locator": "document"}
            count_match = re.search(r"(\d+)\s+source", desc_lower)
            expected = {"source_count": int(count_match.group(1))} if count_match else {"source_count": 1}

        if check_type is None:
            return None

        return {
            "question_number": question_number,
            "description": desc,
            "domain": domain,
            "type": check_type,
            "target": target,
            "expected": expected,
            "marks": marks,
        }

    def _call_openai(self, prompt: str) -> str:
        if self.openai is None:
            self.openai = self._load_openai_module()
        if not self.openai:
            raise RuntimeError("OpenAI SDK is not installed. Install it with 'pip install openai' to enable marksheet parsing.")

        api_key = os.environ.get("OPENAI_API_KEY") or getattr(self.openai, "api_key", None)
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY is not configured.")

        self.openai.api_key = api_key
        response = self.openai.ChatCompletion.create(
            model=self.model,
            messages=[
                {"role": "system", "content": "You convert free-form exam instructions into structured marking JSON."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.0,
            max_tokens=1200,
        )
        return response.choices[0].message.content.strip()

    def _call_ollama(self, prompt: str) -> str:
        try:
            import requests
        except ImportError:
            raise RuntimeError("requests is not installed. Install it with 'pip install requests' to use Ollama backend.")

        url = "http://127.0.0.1:11434/v1/chat/completions"
        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": "You convert free-form exam instructions into structured marking JSON."},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.0,
            "max_tokens": 1200,
        }

        session = requests.Session()
        session.trust_env = False
        last_error = None
        for attempt in range(2):
            try:
                response = session.post(
                    url,
                    json=payload,
                    timeout=120,
                )
                if response.status_code == 404:
                    last_error = RuntimeError(
                        "Ollama API endpoint not found. Ensure Ollama is running and use the /v1/chat/completions endpoint. "
                        f"Received 404 from {url}."
                    )
                    continue
                if response.status_code == 405:
                    last_error = RuntimeError(
                        "Ollama API method not allowed. Ensure you are POSTing to /v1/chat/completions. "
                        f"Received 405 from {url}."
                    )
                    continue
                response.raise_for_status()
                data = response.json()
                if isinstance(data, dict) and data.get("error"):
                    last_error = RuntimeError(f"Ollama API returned an error: {data.get('error')}")
                    continue
                return data["choices"][0]["message"]["content"].strip()
            except (requests.RequestException, ValueError, KeyError) as exc:
                last_error = exc
                continue

        if last_error is not None:
            try:
                return self._call_ollama_cli(prompt)
            except RuntimeError as cli_error:
                raise RuntimeError(
                    f"Ollama HTTP API failed and CLI fallback also failed. HTTP error: {last_error}. CLI error: {cli_error}"
                )

        return self._call_ollama_cli(prompt)

    def _call_ollama_cli(self, prompt: str) -> str:
        command = ["ollama", "run", self.model, "--format", "json"]
        env = os.environ.copy()
        env["OLLAMA_HOST"] = "127.0.0.1:11434"
        try:
            result = subprocess.run(
                command,
                text=True,
                encoding="utf-8",
                errors="replace",
                capture_output=True,
                env=env,
                input=prompt,
                timeout=180,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(f"Ollama CLI timed out: {exc}.")

        if result.returncode != 0:
            raise RuntimeError(
                f"Ollama CLI fallback failed (exit code {result.returncode}).\n"
                f"stdout: {result.stdout}\nstderr: {result.stderr}"
            )

        output = result.stdout.strip()
        if not output:
            raise RuntimeError(
                "Ollama CLI fallback returned no output. "
                "Check that the model is available and the prompt is valid."
            )
        return output

    def _build_prompt(
        self,
        question_paper_text: str,
        marksheet_text: str,
        program: str,
        filename: Optional[str],
    ) -> str:
        file_hint = f" for {filename}" if filename else ""
        return (
            "=== STRICT JSON GENERATOR ===\n"
            "You are a strict JSON generator for exam marking tasks.\n\n"
            "OUTPUT RULES (MANDATORY):\n"
            "- ONLY output valid JSON\n"
            "- NO explanations, NO comments, NO extra text\n"
            "- NO markdown, NO ``` code blocks\n"
            "- Output must start with { and end with }\n"
            "- Each object must have ALL required fields\n\n"
            "TASK: Extract every check from the question paper and memo into structured marking JSON.\n"
            "If the memo contains tables, treat each numbered check row as a separate question entry.\n\n"
            "JSON SCHEMA (REQUIRED FIELDS):\n"
            "{\n"
            "  \"questions\": [\n"
            "    {\n"
            "      \"question_number\": string,\n"
            "      \"description\": string,\n"
            "      \"domain\": string,\n"
            "      \"type\": string,\n"
            "      \"target\": object,\n"
            "      \"expected\": object,\n"
            "      \"marks\": number\n"
            "    }\n"
            "  ]\n"
            "}\n\n"
            "VALID DOMAINS (ONLY USE THESE):\n"
            "- paragraph_formatting\n"
            "- font\n"
            "- table\n"
            "- list\n"
            "- advanced\n"
            "- document\n\n"
            "VALID TYPES BY DOMAIN (ONLY USE THESE):\n"
            "paragraph_formatting: alignment, line_spacing, space_before, space_after, contains_text, border, shading, drop_cap, first_line_indent\n"
            "font: color, size, bold, italic, bold_and_color, underline\n"
            "table: cell_text, cell_alignment, merge_horizontal\n"
            "list: list_style, bullet_char, indent_level, list_paragraph_font\n"
            "advanced: bookmark, bibliography, style_applied\n"
            "document: paper_size, orientation, margins, page_border, watermark, header_text, footer_text, hyphenation, header_alignment, footer_alignment, contains_date\n"
            "object: image_width, image_border, image_crop, smartart, smartart_text, smartart_color\n\n"
            "EXTRACTION RULES:\n"
            "1. Parse actual question/check rows from tables or text.\n"
            "2. Do not use 'type': 'target' or any unsupported type.\n"
            "3. target must be an object with a locator, e.g. {\"locator\": \"after_heading\", \"value\": \"Introduction\"}.\n"
            "4. If you cannot map a check to a supported domain/type, skip it.\n"
            "5. Use the exact question numbering from the memo whenever possible.\n"
            "6. Keep descriptions concise and factual.\n"
            "7. For table-based memos, each row after headers is a question; extract criteria into domain/type/target/expected.\n"
            "8. Map common checks: alignment to paragraph_formatting/alignment, font size to font/size, etc.\n\n"
            "EXAMPLE OUTPUT:\n"
            "{\n"
            "  \"questions\": [\n"
            "    {\n"
            "      \"question_number\": \"1.1\",\n"
            "      \"description\": \"Heading alignment\",\n"
            "      \"domain\": \"paragraph_formatting\",\n"
            "      \"type\": \"alignment\",\n"
            "      \"target\": {\"locator\": \"after_heading\", \"value\": \"Question 1\"},\n"
            "      \"expected\": {\"alignment\": \"center\"},\n"
            "      \"marks\": 2\n"
            "    }\n"
            "  ]\n"
            "}\n\n"
            "INPUT DATA:\n"
            f"Question paper{file_hint}:\n{question_paper_text}\n\n"
            f"Marksheet{file_hint}:\n{marksheet_text}\n\n"
            "OUTPUT: Valid JSON only. Do not output anything else."
        )

    def _extract_json(self, response_text: str) -> List[Dict[str, Any]]:
        for candidate in self._extract_json_candidates(response_text):
            try:
                parsed = json.loads(candidate)
                if isinstance(parsed, dict) and "questions" in parsed:
                    return parsed["questions"]
                if isinstance(parsed, list):
                    return parsed
            except json.JSONDecodeError:
                continue
        raise ValueError("Unable to extract valid JSON from LLM response.")

    def _extract_json_candidates(self, text: str) -> List[str]:
        candidates = []
        if text.strip().startswith("{") or text.strip().startswith("["):
            candidates.append(text)
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1:
            candidates.append(text[start : end + 1])
        start = text.find("[")
        end = text.rfind("]")
        if start != -1 and end != -1:
            candidates.append(text[start : end + 1])
        return candidates

    def _validate_questions(self, questions: List[Dict[str, Any]]) -> List[str]:
        warnings: List[str] = []
        valid_domains = {"paragraph_formatting", "font", "table", "list", "advanced", "document", "object"}
        for idx, question in enumerate(questions, start=1):
            if not isinstance(question, dict):
                warnings.append(f"Question {idx} is not a JSON object.")
                continue
            for field in ["question_number", "description", "domain", "type", "target", "expected", "marks"]:
                if field not in question:
                    warnings.append(f"Question {idx} missing required field: {field}.")
            domain = question.get("domain", "").lower()
            if domain not in valid_domains:
                question["domain"] = "paragraph_formatting"
                warnings.append(f"Question {idx} domain '{domain}' not recognized; using 'paragraph_formatting' as fallback.")
            if not isinstance(question.get("target", {}), dict):
                warnings.append(f"Question {idx} target must be an object.")
        return warnings

