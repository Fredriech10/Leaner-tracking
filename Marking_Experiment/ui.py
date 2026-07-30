"""Tkinter UI for the Marking Experiment prototype."""

from __future__ import annotations

import json
import logging
import subprocess
import threading
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

from .engine import MarkingEngine
from .marking_experiment import MARKING_SCHEMA
from .marksheet_parser import MarksheetParser
from .ollama_manager import list_installed_models
from .task_template import build_task_template, save_task_template

logger = logging.getLogger("marking_experiment_ui")


def _setup_logging() -> None:
    # File logger (keeps logs even if UI shows only a short message).
    # We avoid duplicate handlers if main() is called multiple times.
    if logger.handlers:
        return

    log_path = Path(__file__).resolve().parent / "marking_experiment.log"
    log_path.parent.mkdir(parents=True, exist_ok=True)

    logger.setLevel(logging.INFO)
    file_handler = logging.FileHandler(log_path, encoding="utf-8")
    file_handler.setLevel(logging.INFO)
    formatter = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s")
    file_handler.setFormatter(formatter)

    logger.addHandler(file_handler)


class MarkingExperimentApp:
    def __init__(self, root: tk.Tk) -> None:
        _setup_logging()
        self.root = root
        self.root.title("Marking Experiment")
        self.engine = MarkingEngine()
        self.task_path: Path | None = None
        self.question_paper_path: Path | None = None
        self.marksheet_path: Path | None = None
        self.file_path: Path | None = None
        self.task_definition: dict | None = None
        self.ollama_models: list[str] = []
        self.ollama_ready = False

        self._reset_ollama_on_startup()
        self._build_ui()
        self._update_llm_status()
        self._update_generate_button_state()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

    def _build_ui(self) -> None:
        self.root.columnconfigure(0, weight=1)

        notebook = ttk.Notebook(self.root)
        notebook.grid(row=0, column=0, sticky="nsew")
        self.root.rowconfigure(0, weight=0)

        generate_frame = ttk.Frame(notebook, padding=12)
        marking_frame = ttk.Frame(notebook, padding=12)
        notebook.add(generate_frame, text="Generate Task JSON")
        notebook.add(marking_frame, text="Run Marking")

        generate_frame.columnconfigure(1, weight=1)
        ttk.Label(generate_frame, text="Exam Paper (Word/Excel):").grid(row=0, column=0, sticky="w")
        self.question_paper_entry = ttk.Entry(generate_frame, state="readonly")
        self.question_paper_entry.grid(row=0, column=1, sticky="ew", padx=(8, 8))
        ttk.Button(generate_frame, text="Browse...", command=self._browse_question_paper).grid(row=0, column=2)

        ttk.Label(generate_frame, text="Marksheet / Rubric (Word/Excel):").grid(row=1, column=0, sticky="w", pady=(8, 0))
        self.marksheet_entry = ttk.Entry(generate_frame, state="readonly")
        self.marksheet_entry.grid(row=1, column=1, sticky="ew", padx=(8, 8), pady=(8, 0))
        ttk.Button(generate_frame, text="Browse...", command=self._browse_marksheet).grid(row=1, column=2, pady=(8, 0))

        ttk.Label(generate_frame, text="Ollama Model:").grid(row=2, column=0, sticky="w", pady=(12, 0))
        self.model_var = tk.StringVar(value=self.ollama_models[0] if self.ollama_models else "")
        self.model_combo = ttk.Combobox(generate_frame, textvariable=self.model_var, values=self.ollama_models, state="readonly")
        self.model_combo.grid(row=2, column=1, sticky="ew", padx=(8, 8), pady=(12, 0))
        ttk.Button(generate_frame, text="Refresh", command=self._refresh_ollama_models).grid(row=2, column=2, pady=(12, 0))

        ttk.Label(generate_frame, text="LLM Status:").grid(row=3, column=0, sticky="w", pady=(8, 0))
        self.llm_status_label = ttk.Label(generate_frame, text="Not tested", foreground="red")
        self.llm_status_label.grid(row=3, column=1, sticky="w", padx=(8, 8), pady=(8, 0))
        ttk.Button(generate_frame, text="Test Ollama", command=self._test_llm).grid(row=3, column=2, pady=(8, 0))

        self.llm_response_var = tk.StringVar(value="LLM response will appear here.")
        ttk.Label(generate_frame, textvariable=self.llm_response_var, wraplength=760).grid(row=5, column=1, columnspan=2, sticky="w", padx=(8, 8), pady=(8, 0))

        button_frame = ttk.Frame(generate_frame, padding=(0, 12, 0, 0))
        button_frame.grid(row=6, column=0, columnspan=3, sticky="ew")
        button_frame.columnconfigure(2, weight=1)
        ttk.Button(button_frame, text="Generate Template", command=self._generate_template).grid(row=0, column=0)
        self.generate_button = ttk.Button(button_frame, text="Generate Task JSON", command=self._generate_task_json, state="disabled")
        self.generate_button.grid(row=0, column=1, padx=(8, 0))
        ttk.Label(button_frame, text="Step 1: Upload exam paper and rubric, then generate a JSON task file.").grid(row=0, column=2, sticky="e")

        marking_frame.columnconfigure(1, weight=1)
        ttk.Label(marking_frame, text="Task JSON (generated from exam paper + marksheet):").grid(row=0, column=0, sticky="w")
        self.task_entry = ttk.Entry(marking_frame, state="readonly")
        self.task_entry.grid(row=0, column=1, sticky="ew", padx=(8, 8))
        ttk.Button(marking_frame, text="Browse...", command=self._browse_task).grid(row=0, column=2)

        ttk.Label(marking_frame, text="Learner Submission File:").grid(row=1, column=0, sticky="w", pady=(8, 0))
        self.file_entry = ttk.Entry(marking_frame, state="readonly")
        self.file_entry.grid(row=1, column=1, sticky="ew", padx=(8, 8), pady=(8, 0))
        ttk.Button(marking_frame, text="Browse...", command=self._browse_file).grid(row=1, column=2, pady=(8, 0))

        self.program_label = ttk.Label(marking_frame, text="Program: N/A")
        self.program_label.grid(row=2, column=0, columnspan=2, sticky="w", pady=(12, 0))
        self.question_count_label = ttk.Label(marking_frame, text="Questions: 0")
        self.question_count_label.grid(row=2, column=1, columnspan=2, sticky="e", pady=(12, 0))

        button_frame2 = ttk.Frame(marking_frame, padding=(0, 12, 0, 0))
        button_frame2.grid(row=3, column=0, columnspan=3, sticky="ew")
        button_frame2.columnconfigure(1, weight=1)
        self.run_button = ttk.Button(button_frame2, text="Run Marking", command=self._run_marking, state="disabled")
        self.run_button.grid(row=0, column=0)
        ttk.Label(button_frame2, text="Step 2: Load the generated JSON and the learner's submission, then click Run Marking.").grid(row=0, column=1, sticky="e")

        self.score_label = ttk.Label(marking_frame, text="Score: 0 / 0")
        self.score_label.grid(row=4, column=0, columnspan=3, sticky="w", pady=(12, 0))

        self.tree = ttk.Treeview(
            marking_frame,
            columns=("marks", "passed", "expected", "found"),
            show="tree headings",
            selectmode="browse",
        )
        self.tree.heading("#0", text="Question")
        self.tree.heading("marks", text="Marks")
        self.tree.heading("passed", text="Passed")
        self.tree.heading("expected", text="Expected")
        self.tree.heading("found", text="Found")
        self.tree.column("marks", width=80, anchor="center")
        self.tree.column("passed", width=80, anchor="center")
        self.tree.column("expected", width=180, anchor="w")
        self.tree.column("found", width=180, anchor="w")
        self.tree.grid(row=5, column=0, columnspan=3, sticky="nsew", padx=(0, 0), pady=(8, 0))
        marking_frame.rowconfigure(5, weight=1)

        detail_frame = ttk.Frame(marking_frame, padding=(0, 8, 0, 0))
        detail_frame.grid(row=6, column=0, columnspan=3, sticky="nsew")
        ttk.Label(detail_frame, text="Selected check feedback:").pack(anchor="w")
        self.detail_text = tk.Text(detail_frame, height=6, wrap="word", state="disabled")
        self.detail_text.pack(fill="both", expand=True)
        marking_frame.rowconfigure(6, weight=1)

        self.tree.bind("<<TreeviewSelect>>", self._on_tree_select)

    def _reset_ollama_on_startup(self) -> None:
        self.ollama_models = list_installed_models()

    def _refresh_ollama_models(self) -> None:
        self.ollama_models = list_installed_models()
        self.model_combo.config(values=self.ollama_models)
        if self.ollama_models and not self.model_var.get():
            self.model_var.set(self.ollama_models[0])
        if not self.ollama_models:
            self.model_var.set("")
        self._update_llm_status()

    def _on_close(self) -> None:
        self.root.destroy()

    def _update_llm_status(self) -> None:
        if self.ollama_ready:
            self.llm_status_label.config(text="Ollama ready", foreground="green")
        else:
            self.llm_status_label.config(text="Not tested", foreground="red")
        self._update_generate_button_state()

    def _test_llm(self) -> None:
        model = self.model_var.get().strip()
        if not model:
            self.llm_response_var.set("LLM test failed: no model selected")
            self.llm_status_label.config(text="LLM not ready", foreground="red")
            self.ollama_ready = False
            self._update_generate_button_state()
            return

        parser = MarksheetParser(model=model)
        self.llm_response_var.set("Testing Ollama...")
        self.llm_status_label.config(text="Testing...", foreground="orange")

        def run_test() -> None:
            try:
                def progress_cb(text: str) -> None:
                    # update UI from main thread
                    self.root.after(0, lambda: self.llm_response_var.set(f"LLM response (stream): {text[:800]}"))

                response_text = parser._call_ollama("Hello", progress_cb=progress_cb)
                self.root.after(0, lambda: self.llm_response_var.set(f"LLM response: {response_text}"))
                self.root.after(0, lambda: self.llm_status_label.config(text="Ollama ready", foreground="green"))
                self.ollama_ready = True
                self.root.after(0, self._update_generate_button_state)
            except Exception as exc:
                self.root.after(0, lambda: self.llm_response_var.set(f"LLM test failed: {exc}"))
                self.root.after(0, lambda: self.llm_status_label.config(text="LLM not ready", foreground="red"))
                self.ollama_ready = False
                self.root.after(0, self._update_generate_button_state)

        threading.Thread(target=run_test, daemon=True).start()

    def _update_generate_button_state(self) -> None:
        enabled = bool(self.question_paper_path and self.marksheet_path and self.ollama_ready)
        self.generate_button.config(state="normal" if enabled else "disabled")

    def _browse_task(self) -> None:
        path = filedialog.askopenfilename(
            title="Select task JSON",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialdir=Path(__file__).parent,
        )
        if not path:
            return
        self.task_path = Path(path)
        self.task_entry.config(state="normal")
        self.task_entry.delete(0, tk.END)
        self.task_entry.insert(0, str(self.task_path))
        self.task_entry.config(state="readonly")
        self._load_task_definition()

    def _browse_question_paper(self) -> None:
        path = filedialog.askopenfilename(
            title="Select question paper",
            filetypes=[
                ("Word documents", "*.docx"),
                ("Excel workbooks", "*.xlsx;*.xls"),
                ("Text files", "*.txt"),
                ("All files", "*.*"),
            ],
            initialdir=Path.cwd(),
        )
        if not path:
            return
        self.question_paper_path = Path(path)
        self.question_paper_entry.config(state="normal")
        self.question_paper_entry.delete(0, tk.END)
        self.question_paper_entry.insert(0, str(self.question_paper_path))
        self.question_paper_entry.config(state="readonly")
        self._update_generate_button_state()

    def _browse_marksheet(self) -> None:
        path = filedialog.askopenfilename(
            title="Select marksheet",
            filetypes=[
                ("Word documents", "*.docx"),
                ("Excel workbooks", "*.xlsx;*.xls"),
                ("Text files", "*.txt"),
                ("All files", "*.*"),
            ],
            initialdir=Path.cwd(),
        )
        if not path:
            return
        self.marksheet_path = Path(path)
        self.marksheet_entry.config(state="normal")
        self.marksheet_entry.delete(0, tk.END)
        self.marksheet_entry.insert(0, str(self.marksheet_path))
        self.marksheet_entry.config(state="readonly")
        self._update_generate_button_state()

    def _browse_file(self) -> None:
        path = filedialog.askopenfilename(
            title="Select student file",
            filetypes=[
                ("Word documents", "*.docx"),
                ("Excel workbooks", "*.xlsx"),
                ("HTML files", "*.html;*.htm"),
                ("All files", "*.*"),
            ],
            initialdir=Path.cwd(),
        )
        if not path:
            return
        self.file_path = Path(path)
        self.file_entry.config(state="normal")
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, str(self.file_path))
        self.file_entry.config(state="readonly")
        self._update_run_button_state()

    def _load_task_definition(self) -> None:
        if not self.task_path or not self.task_path.exists():
            return
        try:
            with self.task_path.open("r", encoding="utf-8") as file:
                self.task_definition = json.load(file)
        except Exception as exc:
            messagebox.showerror("Load task", f"Unable to load task JSON:\n{exc}")
            self.task_definition = None
            return

        program = self.task_definition.get("program", "N/A")
        question_count = len(self.task_definition.get("questions", []))
        self.program_label.config(text=f"Program: {program}")
        self.question_count_label.config(text=f"Questions: {question_count}")
        self._populate_task_preview()
        self._update_run_button_state()

    def _populate_task_preview(self) -> None:
        self.tree.delete(*self.tree.get_children())
        if not self.task_definition:
            return
        questions = self.task_definition.get("questions", [])
        for question in questions:
            question_num = question.get("question_number", "?")
            description = question.get("description", "No description")
            marks = question.get("marks", 0)
            domain = question.get("domain", "N/A")
            question_type = question.get("type", "N/A")
            item_id = self.tree.insert(
                "",
                "end",
                text=f"{question_num} - {description[:60]}{'...' if len(description) > 60 else ''}",
                values=(marks, domain, "", ""),
            )
            self.tree.insert(item_id, "end", text=f"Type: {question_type}", values=("", "", "", ""))
            target = question.get("target", {})
            if isinstance(target, dict):
                for key, value in target.items():
                    self.tree.insert(item_id, "end", text=f"Target: {key} = {value}", values=("", "", "", ""))
            expected = question.get("expected", {})
            if isinstance(expected, dict):
                for key, value in expected.items():
                    self.tree.insert(item_id, "end", text=f"Expected: {key} = {value}", values=("", "", "", ""))

    def _generate_template(self) -> None:
        path = filedialog.asksaveasfilename(
            title="Save task JSON template",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*")],
            initialdir=Path.cwd(),
        )
        if not path:
            return
        template = build_task_template()
        save_task_template(Path(path), template)
        messagebox.showinfo("Generate template", f"Template saved to:\n{path}")

    def _generate_task_json(self) -> None:
        if not self.question_paper_path or not self.marksheet_path:
            messagebox.showwarning(
                "Generate task JSON",
                "Please select both a question paper and a marksheet before generating task JSON.",
            )
            return

        logger.info(
            "Generate Task JSON clicked. question_paper=%s marksheet=%s",
            str(self.question_paper_path),
            str(self.marksheet_path),
        )

        ext = self.question_paper_path.suffix.lower()
        if ext in {".xlsx", ".xls"}:
            program = "excel"
        elif ext in {".html", ".htm"}:
            program = "html"
        elif ext in {".accdb", ".mdb"}:
            program = "access"
        else:
            program = "word"

        try:
            question_paper_text = self._read_text_file(self.question_paper_path)
            marksheet_text = self._read_text_file(self.marksheet_path)
        except Exception as exc:
            messagebox.showerror("Generate task JSON", f"Unable to read input files:\n{exc}")
            return

        model = self.model_var.get().strip()
        if not model:
            messagebox.showerror("Generate task JSON", "No Ollama model selected. Install a model with 'ollama pull <model>' and click Refresh.")
            return

        parser = MarksheetParser(model=model)
        # Disable UI while generating
        self.generate_button.config(state="disabled")
        self.llm_status_label.config(text="Generating...", foreground="orange")

        def progress_cb(text: str) -> None:
            # show partial streaming output (trim to reasonable length)
            self.root.after(0, lambda: self.llm_response_var.set(f"LLM stream: {text[:800]}"))

        def run_generate() -> None:
            try:
                parsed = parser.parse(
                    question_paper_text,
                    marksheet_text,
                    program=program,
                    filename=str(self.marksheet_path),
                    progress_cb=progress_cb,
                )
            except Exception as exc:
                logger.exception(
                    "Generate task JSON failed (exception). program=%s question_paper=%s marksheet=%s",
                    program,
                    str(self.question_paper_path),
                    str(self.marksheet_path),
                )
                log_path = Path(__file__).resolve().parent / "marking_experiment.log"
                self.root.after(0, lambda: messagebox.showerror(
                    "Generate task JSON",
                    f"Unable to generate task JSON:\n{exc}\n\nSee logs: {log_path}",
                ))
                self.root.after(0, lambda: self.generate_button.config(state="normal"))
                self.root.after(0, lambda: self.llm_status_label.config(text="LLM not ready", foreground="red"))
                return

            # Success: update UI on main thread
            def on_success() -> None:
                if parsed.warnings:
                    logger.warning(
                        "Generate Task JSON produced warnings. program=%s warnings=%s",
                        program,
                        parsed.warnings,
                    )
                    messagebox.showwarning(
                        "Generate task JSON",
                        "Task JSON generated, but there were warnings:\n" + "\n".join(parsed.warnings) + "\n\nSee logs for details.",
                    )

                default_file = self.file_path.name if self.file_path else f"student_file.{ext.lstrip('.') }"
                # Clamp every generated question to 1 mark.
                split_questions: list[dict] = []

                for q in parsed.questions:
                    q_marks = 1
                    expected = q.get("expected")
                    if isinstance(expected, dict) and len(expected.keys()) > 1:
                        keys = list(expected.keys())
                        for i, k in enumerate(keys):
                            sub_q = {**q}
                            sub_q["question_number"] = f"{sub_q.get('question_number')}{chr(ord('a') + i)}"
                            sub_q["expected"] = {k: expected[k]}
                            sub_q["marks"] = q_marks
                            split_questions.append(sub_q)
                    else:
                        split_q = {**q, "marks": q_marks}
                        split_questions.append(split_q)

                # Prepare task definition and update entries
                task_def = {
                    "task_name": "Generated Marking Task",
                    "program": program,
                    "file": default_file,
                    "total_marks": sum(int(q.get("marks", 1)) for q in split_questions),
                    "questions": split_questions,
                }

                self.task_definition = task_def
                # update task entry
                self.task_entry.config(state="normal")
                # write to a temp file name in UI
                tmp = Path.cwd() / "generated_task.json"
                tmp.write_text(json.dumps(task_def, indent=2, ensure_ascii=False), encoding="utf-8")
                self.task_entry.delete(0, tk.END)
                self.task_entry.insert(0, str(tmp))
                self.task_entry.config(state="readonly")

                self.llm_status_label.config(text="Ollama ready", foreground="green")
                self.generate_button.config(state="normal")

            self.root.after(0, on_success)

        threading.Thread(target=run_generate, daemon=True).start()

        save_path = filedialog.asksaveasfilename(
            title="Save generated task JSON",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialdir=Path.cwd(),
        )
        if save_path:
            with open(save_path, "w", encoding="utf-8") as file:
                json.dump(task_definition, file, indent=2, ensure_ascii=False)
            messagebox.showinfo("Generate task JSON", f"Task JSON saved to:\n{save_path}")
            self.task_path = Path(save_path)
            self.task_entry.config(state="normal")
            self.task_entry.delete(0, tk.END)
            self.task_entry.insert(0, str(self.task_path))
            self.task_entry.config(state="readonly")
            self.task_definition = task_definition
            self._load_task_definition()

    def _read_text_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".txt":
            try:
                return path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                return path.read_text(encoding="cp1252")
        if suffix in {".csv", ".md", ".json", ".yaml", ".yml"}:
            try:
                return path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                return path.read_text(encoding="cp1252")
        if suffix == ".docx":
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            _align_names = {
                WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
                WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
                WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
                WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
            }

            document = Document(path)
            parts = []

            for idx, paragraph in enumerate(document.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                hints = []
                fmt = paragraph.paragraph_format
                align = _align_names.get(paragraph.alignment)
                if align:
                    hints.append(f"align={align}")
                if fmt.line_spacing is not None:
                    hints.append(f"line_spacing={fmt.line_spacing}")
                if fmt.space_before is not None:
                    try:
                        hints.append(f"space_before={fmt.space_before.pt}pt")
                    except Exception:
                        pass
                if fmt.space_after is not None:
                    try:
                        hints.append(f"space_after={fmt.space_after.pt}pt")
                    except Exception:
                        pass
                for run in paragraph.runs:
                    if run.text.strip():
                        if run.font.size:
                            try:
                                hints.append(f"font_size={run.font.size.pt}pt")
                            except Exception:
                                pass
                        if run.font.bold:
                            hints.append("bold")
                        if run.font.italic:
                            hints.append("italic")
                        if run.font.name:
                            hints.append(f"font={run.font.name}")
                        break
                style_name = paragraph.style.name if paragraph.style else ""
                hint_str = f" [{', '.join(hints)}]" if hints else ""
                style_str = f" (style: {style_name})" if style_name else ""
                parts.append(f"[P{idx}]{style_str} {text}{hint_str}")

            for table_idx, table in enumerate(document.tables, start=1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append("\t".join(cells))
                parts.append(f"TABLE {table_idx}:\n" + "\n".join(rows))

            return "\n\n".join(parts)
        if suffix in {".xlsx", ".xls"}:
            try:
                from openpyxl import load_workbook
            except ImportError:
                raise RuntimeError("openpyxl is required to read Excel files. Install it with 'pip install openpyxl'.")

            workbook = load_workbook(path, data_only=True)
            worksheet_texts = []
            for sheet in workbook.worksheets:
                worksheet_texts.append(f"Sheet: {sheet.title}")
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                    rows.append(row_text)
                worksheet_texts.append("\n".join(rows))
            return "\n\n".join(worksheet_texts)
        raise RuntimeError(
            f"Unsupported input file type '{suffix}'. Please use .txt, .docx, .xlsx, .xls, .csv, .md, .json, .yaml or .yml."
        )

    def _update_run_button_state(self) -> None:
        enabled = self.task_definition is not None and self.file_path is not None
        self.run_button.config(state="normal" if enabled else "disabled")

    def _run_marking(self) -> None:
        if not self.task_definition or not self.file_path:
            return
        try:
            session = self.engine.run_task(self.task_definition, self.file_path)
        except Exception as exc:
            messagebox.showerror("Run marking", f"Unable to run marking:\n{exc}")
            return

        # Debug print: show full per-check outcomes in console/log so we can see why checks are skipped.
        try:
            passed = sum(1 for r in getattr(session, 'results', []) if getattr(r, 'passed', False))
            total = len(getattr(session, 'results', []))
            skipped = sum(1 for r in getattr(session, 'results', []) if not getattr(r, 'passed', False))
            print(f"[Marking_Experiment UI] Finished marking. Passed={passed}/{total} Skipped/Failed={skipped}")
            for r in getattr(session, 'results', []):
                print(
                    f"[Marking_Experiment UI] q={r.question_number} marks={r.marks} passed={r.passed} domain/type={getattr(r.details, 'domain', '')}/{getattr(r.details, 'type', '')}"
                )
        except Exception:
            pass

        self._populate_results(session)


    def _populate_results(self, session: object) -> None:
        self.tree.delete(*self.tree.get_children())
        self.score_label.config(text=f"Score: {session.score} / {session.total_marks}")
        for result in session.results:
            expected_text = self._format_result_value(result.expected)
            found_text = self._format_result_value(result.actual)
            item_id = self.tree.insert(
                "",
                "end",
                text=f"{result.question_number} - {result.description}",
                values=(result.marks, "Yes" if result.passed else "No", expected_text, found_text),
            )
            feedback_id = self.tree.insert(
                item_id,
                "end",
                text=f"Feedback: {result.feedback}",
                values=("", "", "", ""),
            )
            self.tree.insert(feedback_id, "end", text=f"Expected: {expected_text}", values=("", "", "", ""))
            self.tree.insert(feedback_id, "end", text=f"Found: {found_text}", values=("", "", "", ""))
            for key, value in result.details.items():
                self.tree.insert(
                    feedback_id,
                    "end",
                    text=f"{key}: {value}",
                    values=("", "", "", ""),
                )

    def _format_result_value(self, value: object) -> str:
        if value is None:
            return "None"
        if isinstance(value, (dict, list, tuple)):
            try:
                return json.dumps(value, ensure_ascii=False, default=str)
            except TypeError:
                return str(value)
        return str(value)

    def _on_tree_select(self, event: tk.Event) -> None:
        selected = self.tree.focus()
        if not selected:
            return
        item = self.tree.item(selected)
        text = item.get("text", "")
        values = item.get("values", [])
        if len(values) >= 4 and (values[2] or values[3]):
            text = f"{text}\n\nExpected: {values[2]}\nFound: {values[3]}"
        self.detail_text.config(state="normal")
        self.detail_text.delete("1.0", tk.END)
        self.detail_text.insert(tk.END, text)
        self.detail_text.config(state="disabled")


def main() -> None:
    root = tk.Tk()
    root.geometry("880x640")
    app = MarkingExperimentApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
