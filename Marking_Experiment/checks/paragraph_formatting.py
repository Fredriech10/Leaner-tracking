"""Paragraph formatting rule checker for Word documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.oxml.ns import qn

from ..checker_types import CheckerResult
from ..utils import compare_numeric, TOLERANCE_CM
from .word_utils import _find_paragraphs, ALIGNMENT_MAP, _read_docx_part
from lxml import etree

NAMESPACES = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def check_paragraph_formatting_rule(rule: Dict[str, Any], file_path: Path) -> CheckerResult:
    document = Document(file_path)
    check_type = str(rule.get("type", ""))
    target = rule.get("target", {}) or {}
    expected = rule.get("expected")

    paragraphs = _find_paragraphs(document, target)
    if not paragraphs:
        return CheckerResult(passed=False, details={"reason": "Paragraph target not found."})
    paragraph = paragraphs[0]
    actual = None
    passed = False

    if check_type == "alignment":
        actual = ALIGNMENT_MAP.get(paragraph.alignment, "left")
        passed = actual == str(expected).lower()
        # If not found on target paragraph, scan all body paragraphs
        if not passed:
            for p in document.paragraphs:
                if ALIGNMENT_MAP.get(p.alignment, "left") == str(expected).lower():
                    passed = True
                    actual = str(expected).lower()
                    break

    elif check_type == "line_spacing":
        # Search all paragraphs for one with explicit line spacing applied
        best_para = paragraph
        for p in document.paragraphs:
            if p.paragraph_format.line_spacing is not None:
                best_para = p
                break
        raw_ls = best_para.paragraph_format.line_spacing
        raw_rule = best_para.paragraph_format.line_spacing_rule
        # Convert EMU to pt (1 pt = 12700 EMU)
        if isinstance(raw_ls, (int, float)) and raw_ls > 100:
            actual_pt = raw_ls / 12700.0
        elif isinstance(raw_ls, (int, float)):
            actual_pt = raw_ls
        else:
            actual_pt = None
        actual = {"rule": str(raw_rule), "value_pt": round(actual_pt, 2) if actual_pt else None}
        if isinstance(expected, dict):
            exp_val = float(expected.get("value", 0))
            exp_unit = expected.get("unit", "pt")
            exp_rule = expected.get("rule", "exact")
            rule_ok = True
            if exp_rule == "exact" and raw_rule is not None:
                rule_ok = "EXACT" in str(raw_rule).upper()
            elif exp_rule == "atLeast" and raw_rule is not None:
                rule_ok = "LEAST" in str(raw_rule).upper()
            if exp_unit == "pt" and actual_pt is not None and exp_val > 0:
                passed = rule_ok and abs(actual_pt - exp_val) < 0.5
            elif exp_unit == "pt" and exp_val == 0:
                passed = rule_ok
            elif exp_unit == "lines" and actual_pt is not None:
                passed = rule_ok and abs(actual_pt - exp_val) < 0.1
            else:
                passed = False
        else:
            passed = str(raw_ls) == str(expected)

    elif check_type == "space_before":
        sb = paragraph.paragraph_format.space_before
        actual = round(sb.pt, 1) if sb else None
        try:
            passed = actual is not None and abs(actual - float(expected)) < 0.5
        except Exception:
            passed = False

    elif check_type == "space_after":
        sa = paragraph.paragraph_format.space_after
        actual = round(sa.pt, 1) if sa else None
        try:
            passed = actual is not None and abs(actual - float(expected)) < 0.5
        except Exception:
            passed = False

    elif check_type == "first_line_indent":
        # Search all paragraphs for one with a first line indent applied
        fi = paragraph.paragraph_format.first_line_indent
        if fi is None:
            for p in document.paragraphs:
                if p.paragraph_format.first_line_indent is not None:
                    fi = p.paragraph_format.first_line_indent
                    break
        actual = float(fi.cm) if fi is not None else None
        try:
            passed = actual is not None and abs(actual - float(expected)) < 0.05
        except Exception:
            passed = False
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "hanging_indent":
        fi = paragraph.paragraph_format.first_line_indent
        actual = abs(float(fi.cm)) if fi is not None and float(fi.cm) < 0 else None
        if actual is not None:
            try:
                passed = abs(actual - float(expected)) < 0.05
            except Exception:
                passed = False
        else:
            passed = False
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "left_indent":
        li = paragraph.paragraph_format.left_indent
        actual = float(li.cm) if li is not None else None
        if actual is not None:
            try:
                passed = abs(actual - float(expected)) < 0.05
            except Exception:
                passed = False
        else:
            passed = False
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "right_indent":
        ri = paragraph.paragraph_format.right_indent
        actual = float(ri.cm) if ri is not None else None
        if actual is not None:
            try:
                passed = abs(actual - float(expected)) < 0.05
            except Exception:
                passed = False
        else:
            passed = False
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "right_to_left":
        actual = _paragraph_has_rtl(paragraph)
        expected_bool = _parse_boolean(expected)
        passed = actual if expected_bool is None else actual == expected_bool
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "keep_with_next":
        actual = bool(paragraph.paragraph_format.keep_with_next)
        expected_bool = _parse_boolean(expected)
        passed = actual if expected_bool is None else actual == expected_bool
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "keep_lines_together":
        actual = bool(paragraph.paragraph_format.keep_together)
        expected_bool = _parse_boolean(expected)
        passed = actual if expected_bool is None else actual == expected_bool
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "widow_orphan_control":
        actual = bool(paragraph.paragraph_format.widow_control)
        expected_bool = _parse_boolean(expected)
        passed = actual if expected_bool is None else actual == expected_bool
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "page_break_before":
        actual = bool(paragraph.paragraph_format.page_break_before)
        expected_bool = _parse_boolean(expected)
        passed = actual if expected_bool is None else actual == expected_bool
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "outline_level":
        actual = _paragraph_outline_level(paragraph)
        if actual is not None:
            exp_val = expected.get("value") if isinstance(expected, dict) else expected
            try:
                passed = int(actual) == int(exp_val)
            except Exception:
                passed = False
        else:
            passed = False
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "tabs":
        passed, actual = _check_tab_stops(paragraph, expected)
        return CheckerResult(passed=passed, actual=actual, details={"type": check_type})

    elif check_type == "contains_text":
        text_to_find = expected.get("text", "") if isinstance(expected, dict) else str(expected)
        found = any(text_to_find.lower() in p.text.lower() for p in document.paragraphs)
        passed = found
        actual = text_to_find if found else None

    elif check_type == "border":
        passed, actual = _check_para_border(paragraph, expected, file_path)

    elif check_type == "shading":
        passed, actual = _check_para_shading(paragraph, expected)

    elif check_type == "drop_cap":
        passed, actual = _check_drop_cap(paragraph, expected)

    else:
        return CheckerResult(passed=False, details={"reason": "Unsupported paragraph formatting check."})

    return CheckerResult(passed=passed, actual=actual, details={"type": check_type})


def _check_para_border(paragraph, expected: Any, file_path: Path) -> tuple[bool, Any]:
    """Check paragraph border via raw XML."""
    COLOUR_MAP = {
        "blue": ["0070c0", "0000ff", "4472c4", "1f3864"],
        "red": ["ff0000", "c00000"],
        "green": ["00b050", "008000"],
        "yellow": ["ffff00", "ffc000"],
        "black": ["000000"],
        "white": ["ffffff"],
    }
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        p_xml = etree.fromstring(paragraph._p.xml.encode("utf-8"))
    except Exception:
        return False, None
    pBdr = p_xml.find(f".//{{{WNS}}}pBdr")
    if pBdr is None:
        return False, {"reason": "No paragraph border found"}
    sides = ["top", "bottom", "left", "right"]
    actual = {}
    for side in sides:
        el = pBdr.find(f"{{{WNS}}}{side}")
        if el is not None:
            actual[side] = {
                "style": el.get(f"{{{WNS}}}val", ""),
                "color": el.get(f"{{{WNS}}}color", "").lower(),
                "sz": el.get(f"{{{WNS}}}sz", ""),  # in 1/8 pt
            }
    if not actual:
        return False, {"reason": "No border sides found"}
    passed = True
    if isinstance(expected, dict):
        exp_style = expected.get("style")
        exp_color = expected.get("color", "").lower() if expected.get("color") else None
        exp_width = expected.get("width_pt")
        for side_data in actual.values():
            if exp_style and exp_style.lower() not in side_data["style"].lower():
                passed = False
            if exp_color:
                hex_val = side_data["color"]
                allowed = COLOUR_MAP.get(exp_color, [exp_color])
                if hex_val not in allowed:
                    passed = False
            if exp_width:
                try:
                    actual_pt = int(side_data["sz"]) / 8.0
                    if abs(actual_pt - exp_width) > 0.5:
                        passed = False
                except Exception:
                    pass
    return passed, actual


def _parse_boolean(value: Any) -> Optional[bool]:
    if isinstance(value, bool):
        return value
    if value is None:
        return None
    normalized = str(value).strip().lower()
    if normalized in {"true", "yes", "1", "on", "y"}:
        return True
    if normalized in {"false", "no", "0", "off", "n"}:
        return False
    return None


def _paragraph_xml_tree(paragraph):
    try:
        return etree.fromstring(paragraph._p.xml.encode("utf-8"))
    except Exception:
        return None


def _paragraph_has_rtl(paragraph) -> bool:
    tree = _paragraph_xml_tree(paragraph)
    if tree is None:
        return False
    bidi = tree.find(".//w:bidi", namespaces=NAMESPACES)
    text_dir = tree.find(".//w:textDirection", namespaces=NAMESPACES)
    return bidi is not None or text_dir is not None


def _paragraph_outline_level(paragraph) -> Optional[int]:
    tree = _paragraph_xml_tree(paragraph)
    if tree is None:
        return None
    outline = tree.find(".//w:outlineLvl", namespaces=NAMESPACES)
    if outline is None:
        return None
    try:
        return int(outline.get(qn("w:val")))
    except Exception:
        return None


def _check_tab_stops(paragraph, expected: Any) -> tuple[bool, Any]:
    tab_stops = paragraph.paragraph_format.tab_stops
    actual = []
    for tab in tab_stops:
        position_cm = None
        try:
            position_cm = round(float(tab.position.cm), 2)
        except Exception:
            pass
        actual.append(
            {
                "position_cm": position_cm,
                "alignment": tab.alignment.name.lower() if hasattr(tab.alignment, "name") else str(tab.alignment).lower() if tab.alignment is not None else None,
                "leader": tab.leader.name.lower() if hasattr(tab.leader, "name") else str(tab.leader).lower() if tab.leader is not None else None,
            }
        )

    if expected is None:
        return (len(actual) > 0, actual)

    passed = True
    if isinstance(expected, dict):
        if expected.get("count") is not None:
            try:
                passed = passed and len(actual) == int(expected["count"])
            except Exception:
                passed = False
        if expected.get("position_cm") is not None:
            position_ok = any(
                compare_numeric(tab["position_cm"], float(expected["position_cm"]), tolerance=TOLERANCE_CM, unit="cm")
                for tab in actual
                if tab["position_cm"] is not None
            )
            passed = passed and position_ok
        if expected.get("alignment") is not None:
            alignment_ok = any(
                tab["alignment"] == str(expected["alignment"]).strip().lower()
                for tab in actual
                if tab["alignment"] is not None
            )
            passed = passed and alignment_ok
        if expected.get("leader") is not None:
            leader_ok = any(
                tab["leader"] == str(expected["leader"]).strip().lower()
                for tab in actual
                if tab["leader"] is not None
            )
            passed = passed and leader_ok
    else:
        try:
            passed = len(actual) == int(expected)
        except Exception:
            passed = False

    return passed, actual


def _check_para_shading(paragraph, expected: Any) -> tuple[bool, Any]:
    """Check paragraph shading/fill via raw XML."""
    SHADING_MAP = {
        "light grey": ["d9d9d9", "bfbfbf", "f2f2f2", "d3d3d3"],
        "light gray": ["d9d9d9", "bfbfbf", "f2f2f2", "d3d3d3"],
        "grey": ["808080", "a5a5a5", "d9d9d9"],
        "gray": ["808080", "a5a5a5", "d9d9d9"],
        "blue": ["0070c0", "0000ff", "4472c4"],
        "yellow": ["ffff00", "ffc000"],
    }
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        p_xml = etree.fromstring(paragraph._p.xml.encode("utf-8"))
    except Exception:
        return False, None
    shd = p_xml.find(f".//{{{WNS}}}shd")
    if shd is None:
        return False, {"reason": "No shading found"}
    fill = shd.get(f"{{{WNS}}}fill", "").lower()
    theme_color = shd.get(f"{{{WNS}}}themeFill", "")
    actual = {"fill": fill, "theme": theme_color}
    if isinstance(expected, dict):
        exp_color = expected.get("color", "any").lower()
        if exp_color == "any":
            passed = bool(fill) and fill != "auto"
        else:
            allowed = SHADING_MAP.get(exp_color, [exp_color])
            passed = fill in allowed or any(a in fill for a in allowed)
    else:
        passed = bool(fill) and fill != "auto"
    return passed, actual


def _check_drop_cap(paragraph, expected: Any) -> tuple[bool, Any]:
    """Check drop cap via raw XML — looks for w:framePr with w:dropCap."""
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        p_xml = etree.fromstring(paragraph._p.xml.encode("utf-8"))
    except Exception:
        return False, None
    frame_pr = p_xml.find(f".//{{{WNS}}}framePr")
    if frame_pr is None:
        return False, {"reason": "No drop cap (framePr) found"}
    drop_cap = frame_pr.get(f"{{{WNS}}}dropCap", "")
    lines = frame_pr.get(f"{{{WNS}}}lines", "")
    actual = {"dropCap": drop_cap, "lines": lines}
    if not drop_cap or drop_cap == "none":
        return False, actual
    passed = True
    if isinstance(expected, dict) and expected.get("lines"):
        try:
            passed = int(lines) == int(expected["lines"])
        except Exception:
            pass
    return passed, actual