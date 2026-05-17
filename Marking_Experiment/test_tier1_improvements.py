"""Integration test for Marking Experiment Tier 1 improvements.

This test verifies:
1. Unit conversions work correctly
2. Numeric comparisons use tolerances
3. Targeting strategies find paragraphs robustly
4. Task validation works
5. WordChecker uses new utilities
"""

from __future__ import annotations

import logging

from Marking_Experiment.config import configure_logging
from Marking_Experiment.utils import (
    pt_to_emu, emu_to_pt, cm_to_emu, emu_to_cm,
    compare_numeric, normalize_hex_color,
    TOLERANCE_PT, TOLERANCE_CM
)
from Marking_Experiment.targeting import string_similarity, Locator
from Marking_Experiment.task_validator import validate_task_definition, constrain_task_to_word

configure_logging(level="DEBUG")


def test_unit_conversions():
    """Test unit conversion accuracy."""
    print("\n=== Testing Unit Conversions ===")
    assert pt_to_emu(12) == 152400, "12 pt should be 152400 EMU"
    assert emu_to_pt(152400) == 12.0, "152400 EMU should be 12.0 pt"
    assert cm_to_emu(2.5) == 900000, "2.5 cm should be 900000 EMU"
    assert emu_to_cm(900000) == 2.5, "900000 EMU should be 2.5 cm"
    print("  [OK] Unit conversions: PASS")


def test_numeric_comparisons():
    """Test numeric comparisons with tolerances."""
    print("\n=== Testing Numeric Comparisons ===")
    assert compare_numeric(12.0, 12.2, tolerance=TOLERANCE_PT, unit="pt") is True
    assert compare_numeric(12.0, 11.8, tolerance=TOLERANCE_PT, unit="pt") is True
    assert compare_numeric(12.0, 12.6, tolerance=TOLERANCE_PT, unit="pt") is False
    assert compare_numeric(2.5, 2.52, tolerance=TOLERANCE_CM, unit="cm") is True
    assert compare_numeric(2.5, 2.57, tolerance=TOLERANCE_CM, unit="cm") is False
    print("  [OK] Numeric comparisons: PASS")


def test_color_normalization():
    """Test hex color normalization."""
    print("\n=== Testing Color Normalization ===")
    assert normalize_hex_color("#FF0000") == "FF0000"
    assert normalize_hex_color("ff0000") == "FF0000"
    assert normalize_hex_color("FF0000") == "FF0000"
    assert normalize_hex_color(None) is None
    assert normalize_hex_color("invalid") is None
    print("  [OK] Color normalization: PASS")


def test_extended_font_checks():
    """Test extended font and character formatting checks."""
    print("\n=== Testing Extended Font Checks ===")
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from Marking_Experiment.checks.font import check_font_rule
    from tempfile import NamedTemporaryFile
    import os

    doc = Document()

    def add_run(text: str, **attrs):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        rPr = run._r.get_or_add_rPr()
        for attr_name, attr_value in attrs.items():
            if attr_name == "dstrike":
                dstrike_el = OxmlElement("w:dstrike")
                rPr.append(dstrike_el)
            elif attr_name == "asciiTheme":
                fonts_el = rPr.find(qn("w:rFonts"))
                if fonts_el is None:
                    fonts_el = OxmlElement("w:rFonts")
                    rPr.append(fonts_el)
                fonts_el.set(qn("w:asciiTheme"), attr_value)
            elif attr_name == "spacing":
                spacing_el = OxmlElement("w:spacing")
                spacing_el.set(qn("w:val"), str(attr_value))
                rPr.append(spacing_el)
            elif attr_name == "kerning":
                kerning_el = OxmlElement("w:kerning")
                kerning_el.set(qn("w:val"), str(attr_value))
                rPr.append(kerning_el)
            else:
                setattr(run.font, attr_name, attr_value)
        return paragraph

    add_run("UNDERLINE", underline=True)
    add_run("STRIKE", strike=True)
    add_run("DOUBLE_STRIKE", dstrike=True)
    add_run("SUPERSCRIPT", superscript=True)
    add_run("SUBSCRIPT", subscript=True)
    add_run("ALLCAPS", all_caps=True)
    add_run("SMALLCAPS", small_caps=True)
    add_run("SHADOW", shadow=True)
    add_run("OUTLINE", outline=True)
    add_run("EMBOSS", emboss=True)
    add_run("HIDDEN", hidden=True)
    add_run("FONTNAME", name="Arial")
    add_run("FONTTHEME", asciiTheme="majorHAnsi")
    add_run("CHAR_SPACING", spacing=120)
    add_run("KERNING", kerning=20)

    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)

        checks = [
            ("underline", "UNDERLINE", True),
            ("strikethrough", "STRIKE", True),
            ("double_strikethrough", "DOUBLE_STRIKE", True),
            ("superscript", "SUPERSCRIPT", True),
            ("subscript", "SUBSCRIPT", True),
            ("all_caps", "ALLCAPS", True),
            ("small_caps", "SMALLCAPS", True),
            ("shadow", "SHADOW", True),
            ("outline", "OUTLINE", True),
            ("emboss", "EMBOSS", True),
            ("hidden", "HIDDEN", True),
            ("font_name", "FONTNAME", "Arial"),
            ("font_theme", "FONTTHEME", "majorHAnsi"),
            ("character_spacing", "CHAR_SPACING", "120"),
            ("kerning", "KERNING", "20"),
        ]

        for check_type, target_text, expected in checks:
            result = check_font_rule(
                {"type": check_type, "target": {"locator": "contains_text", "value": target_text}, "expected": expected},
                tmp_path,
            )
            assert result.passed, f"{check_type} failed: {result}"
    finally:
        os.remove(tmp_path)

    print("  [OK] Extended font checks: PASS")


def test_paragraph_formatting_checks():
    """Test advanced paragraph formatting checks."""
    print("\n=== Testing Paragraph Formatting Checks ===")
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from Marking_Experiment.word_checker import WordChecker
    from tempfile import NamedTemporaryFile
    import os

    doc = Document()
    p_hanging = doc.add_paragraph("HANGING")
    p_hanging.paragraph_format.first_line_indent = Cm(-0.5)
    p_left = doc.add_paragraph("LEFT")
    p_left.paragraph_format.left_indent = Cm(1.0)
    p_right = doc.add_paragraph("RIGHT")
    p_right.paragraph_format.right_indent = Cm(1.0)
    p_keep_next = doc.add_paragraph("KEEP_NEXT")
    p_keep_next.paragraph_format.keep_with_next = True
    p_keep_together = doc.add_paragraph("KEEP_TOGETHER")
    p_keep_together.paragraph_format.keep_together = True
    p_widow = doc.add_paragraph("WIDOW")
    p_widow.paragraph_format.widow_control = True
    p_page_break = doc.add_paragraph("PAGEBREAK")
    p_page_break.paragraph_format.page_break_before = True
    p_outline = doc.add_paragraph("OUTLINE")
    outline_lvl = OxmlElement("w:outlineLvl")
    outline_lvl.set(qn("w:val"), "2")
    p_outline._p.get_or_add_pPr().append(outline_lvl)
    p_tabs = doc.add_paragraph("TABS")
    p_tabs.paragraph_format.tab_stops.add_tab_stop(Cm(2.0))
    p_rtl = doc.add_paragraph("RTL")
    bidi = OxmlElement("w:bidi")
    p_rtl._p.get_or_add_pPr().append(bidi)
    p_align = doc.add_paragraph("ALIGN")
    p_align.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        checker = WordChecker()
        checks = [
            ("hanging_indent", "HANGING", 0.5),
            ("left_indent", "LEFT", 1.0),
            ("right_indent", "RIGHT", 1.0),
            ("keep_with_next", "KEEP_NEXT", True),
            ("keep_lines_together", "KEEP_TOGETHER", True),
            ("widow_orphan_control", "WIDOW", True),
            ("page_break_before", "PAGEBREAK", True),
            ("outline_level", "OUTLINE", 2),
            ("tabs", "TABS", {"count": 1, "position_cm": 2.0}),
            ("right_to_left", "RTL", True),
            ("alignment", "ALIGN", "center"),
        ]
        for check_type, target_text, expected in checks:
            result = checker.check(
                "paragraph_formatting",
                check_type,
                {"locator": "contains_text", "value": target_text},
                expected,
                tmp_path,
            )
            assert result.passed, f"{check_type} failed: {result}"
    finally:
        os.remove(tmp_path)

    print("  [OK] Paragraph formatting checks: PASS")


def test_header_footer_page_layout_checks():
    """Test header/footer content and page layout checks."""
    print("=== Testing Header/Footer and Page Layout Checks ===")
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.shared import Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from Marking_Experiment.word_checker import WordChecker
    from tempfile import NamedTemporaryFile
    import os

    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.start_type = WD_SECTION_START.CONTINUOUS
    section.gutter = Cm(1.0)

    section.header.add_paragraph("STANDARD HEADER")
    first_page_header = section.first_page_header
    first_page_header.add_paragraph("FIRST PAGE HEADER")

    header_page = section.header.add_paragraph()
    header_field = OxmlElement("w:fldSimple")
    header_field.set(qn("w:instr"), " PAGE ")
    header_page._p.append(header_field)

    footer = section.footer
    footer.add_paragraph("FOOTER TEXT")
    footer_field_p = footer.add_paragraph()
    footer_field = OxmlElement("w:fldSimple")
    footer_field.set(qn("w:instr"), " PAGE ")
    footer_field_p._p.append(footer_field)

    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), "FF0000")
    doc._body._element.insert(0, bg)

    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:start"), "3")
    ln.set(qn("w:countBy"), "5")
    section._sectPr.append(ln)

    mirror = OxmlElement("w:mirrorMargins")
    section._sectPr.append(mirror)

    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:fmt"), "lowerRoman")
    section._sectPr.append(pg_num)

    doc.add_page_break()

    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        checker = WordChecker()
        checks = [
            ("header_content", "STANDARD HEADER"),
            ("footer_content", "FOOTER TEXT"),
            ("header_differs", True),
            ("footer_differs", True),
            ("page_number_in_header", True),
            ("page_number_in_footer", True),
            ("page_number_format", "i"),
            ("page_break", True),
            ("section_page_break_type", "continuous"),
            ("line_numbers", {"enabled": True, "start": 3, "interval": 5}),
            ("gutter_margin", 1.0),
            ("mirror_margins", True),
            ("page_color", "FF0000"),
        ]
        for check_type, expected in checks:
            result = checker.check(
                "document",
                check_type,
                {"locator": "document"},
                expected,
                tmp_path,
            )
            assert result.passed, f"{check_type} failed: {result}"
    finally:
        os.remove(tmp_path)

    print("  [OK] Header/Footer and Page Layout Checks: PASS")


def test_string_similarity():
    """Test string similarity for fuzzy matching."""
    print("\n=== Testing String Similarity ===")
    assert string_similarity("Introduction", "Introduction") == 1.0
    assert string_similarity("Introduction", "introduction") == 1.0
    similarity = string_similarity("Introduction", "Introduccion")
    assert 0.85 < similarity < 0.95, f"Typo match should be ~0.92, got {similarity}"
    similarity = string_similarity("Introduction", "Conclusion")
    assert similarity < 0.5, f"Different words should have low similarity, got {similarity}"
    print("  [OK] String similarity: PASS")


def test_task_validation():
    """Test task validation for Word-only support."""
    print("\n=== Testing Task Validation ===")
    valid_task = {
        "task_name": "Test Task",
        "program": "word",
        "total_marks": 5,
        "questions": [
            {
                "question_number": "1",
                "description": "Check alignment",
                "domain": "paragraph_formatting",
                "type": "alignment",
                "target": {"locator": "document"},
                "expected": "left",
                "marks": 1
            }
        ]
    }
    is_valid, warnings = validate_task_definition(valid_task)
    assert is_valid is True, f"Valid task should pass: {warnings}"
    print(f"  [OK] Valid task: PASS (warnings: {len(warnings)})")
    invalid_task = {
        "task_name": "Excel Task",
        "program": "excel",
        "total_marks": 5,
        "questions": []
    }
    is_valid, warnings = validate_task_definition(invalid_task)
    assert is_valid is False, "Excel task should fail validation"
    assert constrain_task_to_word(invalid_task)["program"] == "word"
    print("  [OK] Task validation: PASS")


def test_locator():
    """Test Locator class."""
    print("\n=== Testing Locator ===")
    loc1 = Locator("after_heading", "Introduction")
    assert loc1.kind == "after_heading"
    assert loc1.value == "Introduction"
    loc2 = Locator("paragraph_index", "0")
    assert loc2.kind == "paragraph_index"
    assert loc2.value == "0"
    print("  [OK] Locator: PASS")


def run_all_tests():
    print("\n" + "=" * 60)
    print("MARKING_EXPERIMENT TIER 1 INTEGRATION TESTS")
    print("=" * 60)
    try:
        test_unit_conversions()
        test_numeric_comparisons()
        test_color_normalization()
        test_extended_font_checks()
        test_paragraph_formatting_checks()
        test_string_similarity()
        test_task_validation()
        test_locator()
        print("\n" + "=" * 60)
        print("ALL TESTS PASSED!")
        print("=" * 60)
        return True
    except AssertionError as e:
        print(f"\nTEST FAILED: {e}")
        return False
    except Exception as e:
        print(f"\nERROR: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    import sys
    sys.exit(0 if run_all_tests() else 1)
