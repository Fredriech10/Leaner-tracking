"""Paragraph and table targeting strategies with fallback mechanisms.

This module provides robust locators for finding paragraphs and tables in Word documents.
It includes primary strategies and intelligent fallbacks to handle edge cases.
"""

from __future__ import annotations

import logging
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class Locator:
    """Encapsulates a locator strategy."""
    
    def __init__(self, kind: str, value: Optional[str] = None):
        self.kind = kind  # "paragraph_index", "after_heading", "style_name", etc.
        self.value = value

    def __repr__(self):
        return f"Locator(kind={self.kind}, value={self.value})"


def string_similarity(s1: str, s2: str) -> float:
    """Calculate string similarity (0.0 to 1.0)."""
    if not s1 or not s2:
        return 1.0 if s1 == s2 else 0.0
    return SequenceMatcher(None, s1.lower(), s2.lower()).ratio()


def find_paragraphs_by_index(document: Document, index: int) -> List[Paragraph]:
    """Find paragraph at exact index (primary strategy)."""
    try:
        if 0 <= index < len(document.paragraphs):
            return [document.paragraphs[index]]
    except (IndexError, ValueError):
        pass
    return []


def find_paragraphs_after_heading(document: Document, heading_text: str, fuzzy: bool = True, similarity_threshold: float = 0.8) -> List[Paragraph]:
    """Find paragraph after a heading (primary strategy with fuzzy fallback).
    
    Args:
        document: Word document.
        heading_text: Text to match in heading.
        fuzzy: If True, use fuzzy matching for heading text (handles typos/case differences).
        similarity_threshold: Minimum similarity for fuzzy matching (0.0 to 1.0).
    
    Returns:
        List of next paragraphs after matching heading(s).
    """
    results = []
    heading_lower = heading_text.lower()
    
    for idx, paragraph in enumerate(document.paragraphs):
        para_text = paragraph.text.strip()
        
        # Exact match or substring match (primary)
        if heading_lower in para_text.lower():
            if idx + 1 < len(document.paragraphs):
                next_para = document.paragraphs[idx + 1]
                if next_para.text.strip():  # Skip empty paragraphs
                    results.append(next_para)
        # Fuzzy match (fallback)
        elif fuzzy and similarity_threshold > 0:
            similarity = string_similarity(heading_text, para_text)
            if similarity >= similarity_threshold:
                if idx + 1 < len(document.paragraphs):
                    next_para = document.paragraphs[idx + 1]
                    if next_para.text.strip():
                        results.append(next_para)
                        logger.debug(f"Fuzzy matched heading '{heading_text}' to '{para_text}' (similarity: {similarity:.2f})")
    
    return results


def find_paragraphs_by_style(document: Document, style_name: str, fuzzy: bool = True, similarity_threshold: float = 0.8) -> List[Paragraph]:
    """Find paragraphs with a specific style (primary + fuzzy fallback).
    
    Args:
        document: Word document.
        style_name: Style name to match.
        fuzzy: If True, use fuzzy matching for style names.
        similarity_threshold: Minimum similarity for fuzzy matching.
    
    Returns:
        List of matching paragraphs.
    """
    results = []
    style_lower = style_name.lower()
    
    for paragraph in document.paragraphs:
        if not paragraph.style:
            continue
        
        para_style = paragraph.style.name or ""
        para_style_lower = para_style.lower()
        
        # Exact match (primary)
        if para_style_lower == style_lower or style_lower in para_style_lower:
            results.append(paragraph)
        # Fuzzy match (fallback)
        elif fuzzy and similarity_threshold > 0:
            similarity = string_similarity(style_name, para_style)
            if similarity >= similarity_threshold:
                results.append(paragraph)
                logger.debug(f"Fuzzy matched style '{style_name}' to '{para_style}' (similarity: {similarity:.2f})")
    
    return results


def find_paragraphs_by_text(document: Document, needle: str, exact: bool = False) -> List[Paragraph]:
    """Find paragraphs containing specific text (primary + case-insensitive fallback).
    
    Args:
        document: Word document.
        needle: Text to find.
        exact: If True, find exact (case-sensitive) matches; else case-insensitive.
    
    Returns:
        List of matching paragraphs.
    """
    results = []
    needle_lower = needle.lower() if not exact else needle
    
    for paragraph in document.paragraphs:
        para_text = paragraph.text
        para_text_compare = para_text if exact else para_text.lower()
        
        if needle_lower in para_text_compare:
            results.append(paragraph)
    
    return results


def find_paragraphs_by_style_and_text(document: Document, style_name: str, text_fragment: str) -> List[Paragraph]:
    """Find paragraphs matching both style and text (combined strategy).
    
    Args:
        document: Word document.
        style_name: Style name to match.
        text_fragment: Text fragment to match.
    
    Returns:
        List of matching paragraphs.
    """
    by_style = find_paragraphs_by_style(document, style_name)
    results = [p for p in by_style if text_fragment.lower() in p.text.lower()]
    return results


def find_best_candidate_paragraph(document: Document, target: Dict[str, Any]) -> Optional[Paragraph]:
    """Find the best candidate paragraph using the locator hierarchy.
    
    This function tries multiple strategies in order of reliability:
    1. Exact index
    2. After heading (exact, then fuzzy)
    3. By style name
    4. By text
    5. Combined style + text
    6. First non-empty paragraph (last resort)
    
    Args:
        document: Word document.
        target: Target locator dict with "locator" and "value" keys.
    
    Returns:
        Best matching paragraph, or None if not found.
    """
    locator = str(target.get("locator", "document")).lower()
    value = target.get("value")
    
    candidates = []
    
    # Strategy 1: Exact index (most reliable, if provided)
    if locator == "paragraph_index":
        try:
            idx = int(value)
            candidates = find_paragraphs_by_index(document, idx)
            if candidates:
                logger.debug(f"Found paragraph at index {idx}")
                return candidates[0]
        except (ValueError, TypeError):
            pass
    
    # Strategy 2: After heading (exact, then fuzzy)
    if locator == "after_heading" and value:
        candidates = find_paragraphs_after_heading(document, str(value), fuzzy=True, similarity_threshold=0.75)
        if candidates:
            logger.debug(f"Found {len(candidates)} paragraph(s) after heading '{value}'")
            return candidates[0]
    
    # Strategy 3: By style name
    if locator == "style_name" and value:
        candidates = find_paragraphs_by_style(document, str(value), fuzzy=True)
        if candidates:
            logger.debug(f"Found {len(candidates)} paragraph(s) with style '{value}'")
            return candidates[0]
    
    # Strategy 4: By text content
    if locator == "contains_text" and value:
        candidates = find_paragraphs_by_text(document, str(value))
        if candidates:
            logger.debug(f"Found {len(candidates)} paragraph(s) containing '{value}'")
            return candidates[0]
    
    # Strategy 5: Starting with text
    if locator == "starts_with" and value:
        value_lower = str(value).lower()
        for para in document.paragraphs:
            if para.text.strip().lower().startswith(value_lower):
                logger.debug(f"Found paragraph starting with '{value}'")
                return para
    
    # Strategy 6: Fallback to document (first non-empty)
    if locator == "document" or locator == "near_text":
        for para in document.paragraphs:
            if para.text.strip():
                logger.debug(f"Fallback: using first non-empty paragraph")
                return para
    
    # Last resort: if target indicates "document", return first paragraph
    if not candidates and locator in ("document", ""):
        if document.paragraphs:
            logger.warning(f"Could not find paragraph with locator {locator} = {value}; using first paragraph")
            return document.paragraphs[0]
    
    logger.warning(f"Could not find paragraph with locator {locator} = {value}")
    return None


def find_table(document: Document, target: Dict[str, Any]) -> Optional[Table]:
    """Find table using locator strategy.
    
    Supports:
    - table_index: exact index
    - near_text: find table near text
    - after_heading: find table after heading
    
    Args:
        document: Word document.
        target: Target locator dict.
    
    Returns:
        Matching table, or None if not found.
    """
    locator = str(target.get("locator", "table_index")).lower()
    value = target.get("value")
    
    # Strategy 1: Exact table index
    if locator == "table_index":
        try:
            idx = int(value)
            if 0 <= idx < len(document.tables):
                return document.tables[idx]
        except (ValueError, TypeError, IndexError):
            pass
    
    # Strategy 2: Find table near text
    if locator in ("near_text", "after_heading", "contains_text") and value:
        needle = str(value).lower()
        
        # First, find the paragraph with text
        target_para = None
        for para in document.paragraphs:
            if needle in para.text.lower():
                target_para = para
                break
        
        if target_para:
            # Find the first table after this paragraph
            para_element = target_para._p
            for element in document.element.body:
                if element == para_element:
                    # Start looking from next element
                    continue
                if element.tag == qn("w:tbl"):
                    tbl = element
                    from docx.table import Table
                    return Table(tbl, document)
    
    logger.warning(f"Could not find table with locator {locator} = {value}")
    return None
