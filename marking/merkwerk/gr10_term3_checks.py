from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from decimal import Decimal
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl import load_workbook



def detect_learner_name(folder_name: str) -> str:
    """Retain batch-marker compatibility without depending on the Merk Werk tree."""
    return folder_name


@dataclass
class ResultRow:
    learner_name: str
    learner_folder: str
    question: str
    criterion: str
    maximum_mark: float
    awarded_mark: Any
    result: str
    evidence: str
    file_path: str

    def as_dict(self) -> dict[str, Any]:
        return self.__dict__.copy()


def _row(
    learner_name: str,
    learner_folder: Path,
    question: str,
    criterion: str,
    maximum_mark: float,
    awarded_mark: Any,
    result: str,
    evidence: str,
    file_path: Path | None,
) -> dict[str, Any]:
    return ResultRow(
        learner_name=learner_name,
        learner_folder=str(learner_folder),
        question=question,
        criterion=criterion,
        maximum_mark=maximum_mark,
        awarded_mark=awarded_mark,
        result=result,
        evidence=evidence,
        file_path=str(file_path) if file_path else "",
    ).as_dict()


def mark_learner(folder: Path, rules: dict[str, Any]) -> list[dict[str, Any]]:
    learner_name = detect_learner_name(folder.name)
    expected_files = rules["expected_files"]
    rows: list[dict[str, Any]] = []

    word_path = _find_file(folder, expected_files["word"])
    excel_path = _find_file(folder, expected_files["excel"])
    html_path = _find_file(folder, expected_files["html"])
    image_path = _find_file(folder, expected_files["image"])

    rows.extend(check_word_doc(learner_name, folder, word_path, rules))
    rows.extend(check_excel_book(learner_name, folder, excel_path, rules))
    rows.extend(check_html_page(learner_name, folder, html_path, image_path, rules))
    return rows


def _find_file(folder: Path, filename: str) -> Path | None:
    matches = [path for path in folder.iterdir() if path.is_file() and path.name.lower() == filename.lower()]
    if len(matches) == 1:
        return matches[0]
    return None


def _missing_file_rows(
    learner_name: str, learner_folder: Path, items: list[tuple[str, str, float]]
) -> list[dict[str, Any]]:
    return [
        _row(
            learner_name,
            learner_folder,
            question,
            criterion,
            marks,
            0,
            "FAIL",
            "Required file missing or duplicate candidate files found.",
            None,
        )
        for question, criterion, marks in items
    ]


def check_word_doc(
    learner_name: str, learner_folder: Path, file_path: Path | None, rules: dict[str, Any]
) -> list[dict[str, Any]]:
    rubric = rules["word_checks"]
    if file_path is None:
        return _missing_file_rows(learner_name, learner_folder, [(item["question"], item["criterion"], item["marks"]) for item in rubric])

    try:
        doc = Document(file_path)
    except Exception as exc:
        return [
            _row(learner_name, learner_folder, item["question"], item["criterion"], item["marks"], 0, "FAIL", f"Unreadable DOCX: {exc}", file_path)
            for item in rubric
        ]

    rows: list[dict[str, Any]] = []
    rows.extend(_check_word_margins(learner_name, learner_folder, file_path, doc))
    rows.append(_check_word_page_size(learner_name, learner_folder, file_path, doc))
    rows.append(_check_word_page_color(learner_name, learner_folder, file_path))
    rows.extend(_check_page_numbers(learner_name, learner_folder, file_path, doc))
    rows.extend(_check_heading_styles(learner_name, learner_folder, file_path, doc, rules))
    rows.extend(_check_hyperlink(learner_name, learner_folder, file_path))
    rows.extend(_check_second_paragraph_formatting(learner_name, learner_folder, file_path, doc, rules))
    rows.extend(_check_comment(learner_name, learner_folder, file_path, doc, rules))
    rows.extend(_check_columns(learner_name, learner_folder, file_path))
    return rows


def _binary_row(learner_name: str, learner_folder: Path, question: str, criterion: str, ok: bool, evidence: str, file_path: Path) -> dict[str, Any]:
    return _row(learner_name, learner_folder, question, criterion, 1, 1 if ok else 0, "PASS" if ok else "FAIL", evidence, file_path)


def _check_word_margins(learner_name: str, learner_folder: Path, file_path: Path, doc: Document) -> list[dict[str, Any]]:
    section = doc.sections[0]
    left_cm = round(section.left_margin.cm, 2)
    right_cm = round(section.right_margin.cm, 2)
    evidence = f"Left={left_cm} cm, Right={right_cm} cm"
    return [
        _binary_row(learner_name, learner_folder, "1.1.1.left", "Left margin set to 3 cm", abs(left_cm - 3.0) < 0.05, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.1.1.right", "Right margin set to 2 cm", abs(right_cm - 2.0) < 0.05, evidence, file_path),
    ]


def _check_word_page_size(learner_name: str, learner_folder: Path, file_path: Path, doc: Document) -> dict[str, Any]:
    section = doc.sections[0]
    width_cm = round(section.page_width.cm, 2)
    height_cm = round(section.page_height.cm, 2)
    ok = (abs(width_cm - 21.0) < 0.1 and abs(height_cm - 29.7) < 0.1) or (abs(width_cm - 29.7) < 0.1 and abs(height_cm - 21.0) < 0.1)
    evidence = f"Page size={width_cm} x {height_cm} cm"
    return _row(learner_name, learner_folder, "1.1.2", "Document size set to A4", 1, 1 if ok else 0, "PASS" if ok else "FAIL", evidence, file_path)


def _check_word_page_color(learner_name: str, learner_folder: Path, file_path: Path) -> dict[str, Any]:
    try:
        with ZipFile(file_path) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        background = re.search(r"<w:background\b[^>]*>", document_xml)
        color = re.search(r'w:color="([0-9A-Fa-f]{6}|auto)"', background.group(0)) if background else None
        ok = background is not None
        evidence = f"Page background={background.group(0)}" if background else "No page background element found"
    except Exception as exc:
        ok = False
        evidence = f"Could not inspect page color: {exc}"
    return _row(learner_name, learner_folder, "1.1.3", "Page background present (teacher override)", 1, 1 if ok else 0, "PASS" if ok else "FAIL", evidence, file_path)


def _check_page_numbers(learner_name: str, learner_folder: Path, file_path: Path, doc: Document) -> list[dict[str, Any]]:
    footer_texts: list[str] = []
    field_codes: list[str] = []
    for section in doc.sections:
        visible_parts: list[str] = []
        for node in section.footer._element.xpath(".//w:t | .//w:instrText"):
            text = (node.text or "").strip()
            if not text:
                continue
            if node.tag.endswith("instrText"):
                field_codes.append(text.upper())
            else:
                visible_parts.append(text)
        if visible_parts:
            footer_texts.append(" ".join(visible_parts))
    visible_footer = " / ".join(dict.fromkeys(footer_texts))
    has_page = any("PAGE" in code and "NUMPAGES" not in code for code in field_codes)
    has_total = any("NUMPAGES" in code for code in field_codes)
    footer_is_page_2_of_2 = visible_footer == "Page 2 of 2"
    return [
        _binary_row(learner_name, learner_folder, "1.2.page", "Footer displays 'Page 2 of 2'", footer_is_page_2_of_2, f"footer={visible_footer!r}", file_path),
        _binary_row(learner_name, learner_folder, "1.2.position", "Footer displays 'Page 2 of 2'", footer_is_page_2_of_2, f"footer={visible_footer!r}", file_path),
        _binary_row(learner_name, learner_folder, "1.2.format", "Footer contains PAGE and NUMPAGES fields", has_page and has_total, f"fields={', '.join(dict.fromkeys(field_codes)) or '(none)'}", file_path),
    ]


def _check_heading_styles(
    learner_name: str, learner_folder: Path, file_path: Path, doc: Document, rules: dict[str, Any]
) -> list[dict[str, Any]]:
    expected = set(rules["word_expected"]["animal_headings"])
    matched = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in expected and paragraph.style and paragraph.style.name == "Heading 2":
            matched.append(paragraph.text.strip())
    count = len(set(matched))
    evidence = f"Heading 2 applied to {count}/{len(expected)} expected animal headings"
    return [
        _binary_row(learner_name, learner_folder, "1.3.style", "Heading 2 style applied to animal headings", count > 0, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.3.all", "Heading 2 style applied to at least 13 animal-name headings", count >= 13, evidence, file_path),
    ]


def _check_hyperlink(learner_name: str, learner_folder: Path, file_path: Path) -> list[dict[str, Any]]:
    try:
        with ZipFile(file_path) as archive:
            document_xml = ET.fromstring(archive.read("word/document.xml"))
            rels_xml = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
        }
        targets = {
            relationship.attrib["Id"]: relationship.attrib.get("Target", "")
            for relationship in rels_xml.findall("pr:Relationship", namespaces)
        }
        links = []
        for hyperlink in document_xml.findall(".//w:hyperlink", namespaces):
            linked_text = "".join(node.text or "" for node in hyperlink.findall(".//w:t", namespaces)).strip()
            relation_id = hyperlink.attrib.get(f"{{{namespaces['r']}}}id")
            links.append((linked_text, targets.get(relation_id, "")))
        hyperlink_present = bool(links)
        target_present = any(target.lower().endswith("1bio_data.docx") for _, target in links)
        biodiversity_link = any(text.lower() == "biodiversity" and target.lower().endswith("1bio_data.docx") for text, target in links)
        evidence = "links=" + (", ".join(f"{text!r}->{target}" for text, target in links) or "(none)")
    except Exception as exc:
        evidence = f"Could not inspect hyperlink: {exc}"
        hyperlink_present = target_present = biodiversity_link = False
    return [
        _binary_row(learner_name, learner_folder, "1.4.hyperlink", "Hyperlink inserted", hyperlink_present, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.4.first_paragraph", "Hyperlink text is 'biodiversity' and targets 1Bio_Data.docx", biodiversity_link, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.4.target", "Hyperlink target is 1Bio_Data.docx", target_present, evidence, file_path),
    ]


def _check_second_paragraph_formatting(
    learner_name: str, learner_folder: Path, file_path: Path, doc: Document, rules: dict[str, Any]
) -> list[dict[str, Any]]:
    marker = rules["word_expected"]["second_paragraph_startswith"]
    paragraph = next((p for p in doc.paragraphs if p.text.strip().startswith(marker)), None)
    if paragraph is None:
        fail = "Target paragraph not found."
        return [
            _row(learner_name, learner_folder, "1.5.1", "Second paragraph justified", 1, 0, "FAIL", fail, file_path),
            _binary_row(learner_name, learner_folder, "1.5.2.spacing", "Second paragraph has 1.5 line spacing", False, fail, file_path),
            _binary_row(learner_name, learner_folder, "1.5.2.after", "Second paragraph has 6 pt spacing after", False, fail, file_path),
        ]
    alignment_ok = paragraph.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    line_spacing = paragraph.paragraph_format.line_spacing
    space_after = _effective_space_after(paragraph, file_path)
    spacing_ok = float(line_spacing or 0) == 1.5 and round(space_after, 1) == 6.0
    return [
        _row(
            learner_name,
            learner_folder,
            "1.5.1",
            "Second paragraph justified",
            1,
            1 if alignment_ok else 0,
            "PASS" if alignment_ok else "FAIL",
            f"Alignment={paragraph.paragraph_format.alignment}",
            file_path,
        ),
        _binary_row(learner_name, learner_folder, "1.5.2.spacing", "Second paragraph has 1.5 line spacing", float(line_spacing or 0) == 1.5, f"Line spacing={line_spacing}", file_path),
        _binary_row(learner_name, learner_folder, "1.5.2.after", "Second paragraph has 6 pt spacing after", round(space_after, 1) == 6.0, f"Space after={round(space_after, 1)} pt", file_path),
    ]


def _effective_space_after(paragraph: Any, file_path: Path) -> float:
    direct = paragraph.paragraph_format.space_after
    if direct is not None:
        return direct.pt
    style = paragraph.style
    while style is not None:
        style_value = style.paragraph_format.space_after
        if style_value is not None:
            return style_value.pt
        style = style.base_style
    try:
        with ZipFile(file_path) as archive:
            styles_xml = archive.read("word/styles.xml").decode("utf-8", errors="ignore")
        match = re.search(r"<w:pPrDefault>.*?<w:spacing\b[^>]*w:after=\"(\d+)\"", styles_xml, flags=re.DOTALL)
        if match:
            return int(match.group(1)) / 20
    except Exception:
        pass
    return 0.0


def _check_comment(
    learner_name: str, learner_folder: Path, file_path: Path, doc: Document, rules: dict[str, Any]
) -> list[dict[str, Any]]:
    expected = rules["word_expected"]["comment_text"].lower()
    try:
        with ZipFile(file_path) as archive:
            names = {entry.filename for entry in archive.infolist()}
            if "word/comments.xml" not in names:
                return [
                    _binary_row(learner_name, learner_folder, "1.6.phrase", "Comment is added to the Platypus 'Where it is found:' phrase", False, "No comments.xml part found.", file_path),
                    _binary_row(learner_name, learner_folder, "1.6.text", "Comment text is 'Animal sections'", False, "No comments.xml part found.", file_path),
                ]
            comments_xml = archive.read("word/comments.xml").decode("utf-8", errors="ignore").lower()
        has_comment = bool(re.search(r"<w:comment(?:\s|>)", comments_xml))
        phrase_ok = any(paragraph.text.strip().lower() == "where it is found:" for paragraph in doc.paragraphs)
        evidence = f"Target phrase found={phrase_ok}; comment found={has_comment}"
    except Exception as exc:
        has_comment = phrase_ok = False
        evidence = f"Could not inspect comments: {exc}"
    return [
        _binary_row(learner_name, learner_folder, "1.6.phrase", "Comment present (teacher override)", has_comment, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.6.text", "Comment present (teacher override)", has_comment, evidence, file_path),
    ]


def _check_columns(learner_name: str, learner_folder: Path, file_path: Path) -> list[dict[str, Any]]:
    try:
        with ZipFile(file_path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore").lower()
        two_columns = bool(re.search(r"<w:cols[^>]*w:num=\"2\"", xml))
        line_between = bool(re.search(r"<w:cols[^>]*w:sep=\"1\"", xml))
        column_break = "<w:br w:type=\"column\"" in xml
        evidence = f"two_columns={two_columns}, line_between={line_between}, column_break={column_break}"
    except Exception as exc:
        two_columns = line_between = column_break = False
        evidence = f"Could not inspect columns: {exc}"
    return [
        _binary_row(learner_name, learner_folder, "1.7.columns", "Narwhal and Quokka section is set to two columns", two_columns, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.7.line", "Line between the two columns", line_between, evidence, file_path),
        _binary_row(learner_name, learner_folder, "1.7.break", "Column break added", column_break, evidence, file_path),
    ]


def _manual_review(
    learner_name: str,
    learner_folder: Path,
    question: str,
    criterion: str,
    marks: float,
    file_path: Path,
    reason: str,
) -> dict[str, Any]:
    return _row(learner_name, learner_folder, question, criterion, marks, "MANUAL REVIEW REQUIRED", "MANUAL REVIEW REQUIRED", reason, file_path)


def check_excel_book(
    learner_name: str, learner_folder: Path, file_path: Path | None, rules: dict[str, Any]
) -> list[dict[str, Any]]:
    rubric = rules["excel_checks"]
    if file_path is None:
        return _missing_file_rows(learner_name, learner_folder, [(item["question"], item["criterion"], item["marks"]) for item in rubric])
    try:
        workbook = load_workbook(file_path, data_only=False)
    except Exception as exc:
        return [
            _row(learner_name, learner_folder, item["question"], item["criterion"], item["marks"], 0, "FAIL", f"Unreadable workbook: {exc}", file_path)
            for item in rubric
        ]

    rows: list[dict[str, Any]] = []
    sheetnames_lower = [name.lower() for name in workbook.sheetnames]
    rows.append(_row(learner_name, learner_folder, "2.1.1", "Sheet1 renamed to More Animals", 1, 1 if "more animals" in sheetnames_lower else 0, "PASS" if "more animals" in sheetnames_lower else "FAIL", f"Sheets={workbook.sheetnames}", file_path))

    sheet2 = workbook["Sheet2"] if "Sheet2" in workbook.sheetnames else None
    rows.extend(_check_sheet2_sort(learner_name, learner_folder, file_path, sheet2))

    animals = workbook["Animals"] if "Animals" in workbook.sheetnames else None
    if animals is None:
        for item in rubric:
            if item["question"].startswith(("2.3", "2.4", "2.5", "2.6", "2.7", "2.8", "2.9", "2.10")):
                rows.append(_row(learner_name, learner_folder, item["question"], item["criterion"], item["marks"], 0, "FAIL", "Animals worksheet missing.", file_path))
        return rows

    rows.extend(_check_formula_components(learner_name, learner_folder, file_path, animals, "B63", "2.3", "AVERAGE", "F2:F61"))
    rows.extend(_check_formula_components(learner_name, learner_folder, file_path, animals, "B64", "2.4", "MAX", "H2:H61"))
    rows.extend(_check_formula_components(learner_name, learner_folder, file_path, animals, "B65", "2.5", "COUNT", "A2:A61", allow_counta=True))
    rows.extend(_check_mode_components(learner_name, learner_folder, file_path, animals))
    rows.extend(_check_formula_components(learner_name, learner_folder, file_path, animals, "B67", "2.7", "COUNTBLANK", "G2:G61"))
    b68 = _normalize_formula(animals["B68"].value)
    rows.append(_binary_row(learner_name, learner_folder, "2.8", "Population difference formula for Narwhal and Red Panda", b68 == "G4-G8", f"B68 formula={animals['B68'].value!r}", file_path))
    b69 = _normalize_formula(animals["B69"].value)
    rows.append(_binary_row(learner_name, learner_folder, "2.9", "Today's date function", b69 == "TODAY()", f"B69 formula={animals['B69'].value!r}", file_path))
    rows.extend(_check_countif_components(learner_name, learner_folder, file_path, animals))
    return rows


def _check_sheet2_sort(learner_name: str, learner_folder: Path, file_path: Path, sheet2: Any) -> list[dict[str, Any]]:
    if sheet2 is None:
        return [
            _binary_row(learner_name, learner_folder, "2.1.2.order", "Sheet2 sorted in A-Z order", False, "Sheet2 missing.", file_path),
            _binary_row(learner_name, learner_folder, "2.1.2.status", "Sort is based on Conservation Status", False, "Sheet2 missing.", file_path),
        ]
    values = [sheet2[f"A{i}"].value for i in range(2, sheet2.max_row + 1) if sheet2[f"A{i}"].value]
    normalized = [str(v) for v in values]
    order_ok = normalized == sorted(normalized, key=str.casefold)
    header = str(sheet2["A1"].value or "").strip().lower()
    status_ok = "conservation" in header and "status" in header
    evidence = f"Observed order={normalized}"
    return [
        _binary_row(learner_name, learner_folder, "2.1.2.order", "Sheet2 sorted in A-Z order", order_ok, evidence, file_path),
        _binary_row(learner_name, learner_folder, "2.1.2.status", "Sort is based on Conservation Status", status_ok, f"Sheet2 A1={sheet2['A1'].value!r}", file_path),
    ]


def _normalize_formula(value: Any) -> str:
    if value is None:
        return ""
    if hasattr(value, "text"):
        value = value.text
    text = str(value)
    if text.startswith("="):
        text = text[1:]
    text = text.replace("$", "").replace(" ", "").upper()
    return text


def _check_formula_components(
    learner_name: str,
    learner_folder: Path,
    file_path: Path,
    sheet: Any,
    cell_ref: str,
    question: str,
    function: str,
    expected_range: str,
    allow_counta: bool = False,
) -> list[dict[str, Any]]:
    cell = sheet[cell_ref]
    formula = _normalize_formula(cell.value)
    function_ok = function in formula and (not allow_counta or ("COUNT(" in formula or "COUNTA(" in formula))
    range_ok = expected_range.upper() in formula
    range_criterion = f"{cell_ref} references {expected_range}"
    if question == "2.5":
        # Teacher rule: accept any count range starting at row 2 and ending at row 61 or 62.
        range_ok = "2" in formula and ("61" in formula or "62" in formula)
        range_criterion = f"{cell_ref} range includes row 2 and row 61 or 62"
    evidence = f"{cell_ref} formula={cell.value!r}"
    return [
        _binary_row(learner_name, learner_folder, f"{question}.function", f"{cell_ref} uses {function}{' or COUNTA' if allow_counta else ''}", function_ok, evidence, file_path),
        _binary_row(learner_name, learner_folder, f"{question}.range", range_criterion, range_ok, evidence, file_path),
    ]


def _check_mode_components(learner_name: str, learner_folder: Path, file_path: Path, sheet: Any) -> list[dict[str, Any]]:
    cell = sheet["B66"]
    formula = _normalize_formula(cell.value)
    evidence = f"B66 formula={cell.value!r}; number_format={cell.number_format!r}"
    return [
        _binary_row(learner_name, learner_folder, "2.6.function", "B66 uses a MODE function", formula.startswith("MODE(" ) or formula.startswith("MODE.SNGL("), evidence, file_path),
        _binary_row(learner_name, learner_folder, "2.6.range", "B66 references D2:D61", "D2:D61" in formula, evidence, file_path),
        _binary_row(learner_name, learner_folder, "2.6.format", "B66 result displays two decimal places", "0.00" in str(cell.number_format), evidence, file_path),
    ]


def _check_countif_components(learner_name: str, learner_folder: Path, file_path: Path, sheet: Any) -> list[dict[str, Any]]:
    formula = _normalize_formula(sheet["B70"].value)
    evidence = f"B70 formula={sheet['B70'].value!r}"
    return [
        _binary_row(learner_name, learner_folder, "2.10.function", "B70 uses COUNTIF", formula.startswith("COUNTIF("), evidence, file_path),
        _binary_row(learner_name, learner_folder, "2.10.range", "B70 references D2:D61", "D2:D61" in formula, evidence, file_path),
        _binary_row(learner_name, learner_folder, "2.10.criteria", "B70 uses a greater-than-or-equal criterion", ">=" in formula, evidence, file_path),
        _binary_row(learner_name, learner_folder, "2.10.value", "B70 criterion uses 40", "40" in formula, evidence, file_path),
    ]


def check_html_page(
    learner_name: str, learner_folder: Path, file_path: Path | None, image_path: Path | None, rules: dict[str, Any]
) -> list[dict[str, Any]]:
    rubric = rules["html_checks"]
    if file_path is None:
        return _missing_file_rows(learner_name, learner_folder, [(item["question"], item["criterion"], item["marks"]) for item in rubric])
    try:
        html = file_path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        return [
            _row(learner_name, learner_folder, item["question"], item["criterion"], item["marks"], 0, "FAIL", f"Unreadable HTML: {exc}", file_path)
            for item in rubric
        ]

    soup = BeautifulSoup(html, "html.parser")
    rows: list[dict[str, Any]] = []

    body = soup.find("body")
    bgcolor = (body.get("bgcolor", "") if body else "").strip().lower()
    rows.append(_row(learner_name, learner_folder, "3.1", "Body background colour set to yellow", 1, 1 if bgcolor == "yellow" else 0, "PASS" if bgcolor == "yellow" else "FAIL", f"bgcolor={bgcolor!r}", file_path))

    font = soup.find("font")
    font_present = False
    face_ok = size_ok = color_ok = False
    if font and "unique animals of the world" in font.get_text(" ", strip=True).lower():
        font_present = True
        face = (font.get("face", "") or "").lower()
        size = str(font.get("size", "")).strip()
        color = (font.get("color", "") or "").lower()
        face_ok = face == "arial"
        size_ok = size == "6"
        color_ok = color == "green"
        evidence = f"face={face}, size={size}, color={color}"
    else:
        evidence = "Main heading font tag not found."
    rows.extend([
        _binary_row(learner_name, learner_folder, "3.2.font", "Main heading uses a font tag", font_present, evidence, file_path),
        _binary_row(learner_name, learner_folder, "3.2.face", "Main-heading font face is Arial", face_ok, evidence, file_path),
        _binary_row(learner_name, learner_folder, "3.2.size", "Main-heading font size is 6", size_ok, evidence, file_path),
        _binary_row(learner_name, learner_folder, "3.2.color", "Main-heading font colour is green", color_ok, evidence, file_path),
    ])

    hrs = soup.find_all("hr")
    width_ok = any((hr.get("width", "") or "").replace(" ", "") == "80%" for hr in hrs)
    rows.append(_row(learner_name, learner_folder, "3.3", "Horizontal line below heading width 80%", 1, 1 if width_ok else 0, "PASS" if width_ok else "FAIL", f"hr_count={len(hrs)}", file_path))

    h3s = soup.find_all("h3")
    italic_ok = bool(h3s) and all(h3.find("i") or h3.find("em") for h3 in h3s)
    rows.extend([
        _binary_row(learner_name, learner_folder, "3.4.h3", "Animal section headings use h3 tags", bool(h3s), f"h3_count={len(h3s)}", file_path),
        _binary_row(learner_name, learner_folder, "3.4.italics", "All h3 section headings are italicised", italic_ok, f"h3_count={len(h3s)}", file_path),
    ])

    paragraphs = soup.find_all("p")
    trailing_hr_ok = False
    if paragraphs and hrs:
        last_p = paragraphs[-1]
        next_tag = last_p.find_next(lambda tag: getattr(tag, "name", None) is not None)
        trailing_hr_ok = next_tag is not None and next_tag.name == "hr"
    rows.append(_row(learner_name, learner_folder, "3.5", "Second horizontal rule after last paragraph", 1, 1 if trailing_hr_ok else 0, "PASS" if trailing_hr_ok else "FAIL", f"trailing_hr={trailing_hr_ok}", file_path))

    nesting_ok = _has_balanced_html_tags(html)
    rows.append(_row(learner_name, learner_folder, "3.closing", "Closing tags and nesting", 1, 1 if nesting_ok else 0, "PASS" if nesting_ok else "FAIL", "Tags are balanced and correctly nested" if nesting_ok else "Unclosed or incorrectly nested HTML tag found", file_path))

    if image_path is None:
        rows.append(_row(learner_name, learner_folder, "3.asset", "Referenced image asset present", 0, 0, "INFO", "3Animals.png missing or duplicated.", file_path))
    return rows


def _has_balanced_html_tags(html: str) -> bool:
    void_tags = {"area", "base", "br", "col", "embed", "hr", "img", "input", "link", "meta", "param", "source", "track", "wbr"}
    stack: list[str] = []
    for match in re.finditer(r"</?([a-zA-Z][\w:-]*)(?:\s[^<>]*)?/?>", html):
        tag = match.group(1).lower()
        token = match.group(0)
        if token.startswith("</"):
            if not stack or stack.pop() != tag:
                return False
        elif tag not in void_tags and not token.rstrip().endswith("/>"):
            stack.append(tag)
    return not stack
