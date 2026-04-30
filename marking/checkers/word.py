"""
checkers/word.py
Reusable check functions for Word (.docx) assignments.
All functions take filepath as first argument.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


# ── Helper ────────────────────────────────────────────────────────────────────

def _load(filepath):
    return Document(filepath)


def _normalize(value):
    if value is None:
        return ""
    return str(value).strip().upper()


# ── Text checks ───────────────────────────────────────────────────────────────

def contains_text(filepath, expected, case_sensitive=False):
    """Check if the document contains a specific string anywhere in the body."""
    doc = _load(filepath)
    full_text = " ".join(p.text for p in doc.paragraphs)
    if case_sensitive:
        return expected in full_text
    return expected.upper() in full_text.upper()


def paragraph_text(filepath, index, expected, exact=False):
    """
    Check the text of a specific paragraph by index (0-based).
    Set exact=True for exact match, otherwise normalized comparison.
    """
    doc = _load(filepath)
    if index >= len(doc.paragraphs):
        return False
    actual = doc.paragraphs[index].text
    if exact:
        return actual == expected
    return _normalize(actual) == _normalize(expected)


def paragraph_count_at_least(filepath, minimum):
    """Check if the document has at least a minimum number of paragraphs."""
    doc = _load(filepath)
    return len(doc.paragraphs) >= minimum


def word_count_at_least(filepath, minimum):
    """Check if the document body has at least a minimum number of words."""
    doc = _load(filepath)
    words = " ".join(p.text for p in doc.paragraphs).split()
    return len(words) >= minimum


# ── Heading checks ────────────────────────────────────────────────────────────

def heading_exists(filepath, level, expected_text=None):
    """
    Check if a heading of a given level exists.
    Optionally check that it contains expected_text.
    Level 1 = Heading 1, Level 2 = Heading 2 etc.
    """
    doc = _load(filepath)
    style_name = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style.name == style_name:
            if expected_text is None:
                return True
            if expected_text.upper() in p.text.upper():
                return True
    return False


# ── Style / formatting checks ─────────────────────────────────────────────────

def paragraph_is_bold(filepath, index):
    """Check if all runs in a paragraph are bold."""
    doc = _load(filepath)
    if index >= len(doc.paragraphs):
        return False
    runs = doc.paragraphs[index].runs
    return all(run.bold for run in runs) if runs else False


def paragraph_font_size(filepath, index, expected_size):
    """
    Check if the first run of a paragraph has a specific font size (in pt).
    expected_size should be an integer e.g. 12
    """
    doc = _load(filepath)
    if index >= len(doc.paragraphs):
        return False
    runs = doc.paragraphs[index].runs
    if not runs:
        return False
    size = runs[0].font.size
    if size is None:
        return False
    return int(size.pt) == expected_size


def paragraph_font_name(filepath, index, expected_font):
    """Check if the first run of a paragraph uses a specific font."""
    doc = _load(filepath)
    if index >= len(doc.paragraphs):
        return False
    runs = doc.paragraphs[index].runs
    if not runs:
        return False
    return (runs[0].font.name or "").upper() == expected_font.upper()


def paragraph_alignment(filepath, index, expected):
    """
    Check paragraph alignment.
    expected: 'left', 'center', 'right', 'justify'
    """
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    doc = _load(filepath)
    if index >= len(doc.paragraphs):
        return False
    return doc.paragraphs[index].alignment == mapping.get(expected.lower())


# ── Table checks ──────────────────────────────────────────────────────────────

def table_exists(filepath):
    """Check if at least one table exists in the document."""
    doc = _load(filepath)
    return len(doc.tables) > 0


def table_count(filepath, expected_count):
    """Check if the document has an exact number of tables."""
    doc = _load(filepath)
    return len(doc.tables) == expected_count


def table_has_rows(filepath, table_index, expected_rows):
    """Check if a specific table has an expected number of rows."""
    doc = _load(filepath)
    if table_index >= len(doc.tables):
        return False
    return len(doc.tables[table_index].rows) == expected_rows


def table_has_columns(filepath, table_index, expected_cols):
    """Check if a specific table has an expected number of columns."""
    doc = _load(filepath)
    if table_index >= len(doc.tables):
        return False
    return len(doc.tables[table_index].columns) == expected_cols


def table_cell_value(filepath, table_index, row, col, expected):
    """Check the text value of a specific table cell (0-based row/col)."""
    doc = _load(filepath)
    if table_index >= len(doc.tables):
        return False
    try:
        actual = doc.tables[table_index].cell(row, col).text
        return _normalize(actual) == _normalize(expected)
    except Exception:
        return False


# ── Page setup checks ─────────────────────────────────────────────────────────

def page_is_landscape(filepath):
    """Check if the first section is landscape orientation."""
    doc = _load(filepath)
    section = doc.sections[0]
    return section.orientation == WD_ORIENT.LANDSCAPE


def page_size_is_a4(filepath):
    """Check if the first section is A4 size (within 1mm tolerance)."""
    from docx.shared import Mm
    doc = _load(filepath)
    section = doc.sections[0]
    a4_width = Mm(210)
    a4_height = Mm(297)
    tolerance = Mm(1)
    width_ok = abs(section.page_width - a4_width) < tolerance
    height_ok = abs(section.page_height - a4_height) < tolerance
    return width_ok and height_ok


# ── Image / object checks ─────────────────────────────────────────────────────

def image_exists(filepath):
    """Check if the document contains at least one embedded image."""
    doc = _load(filepath)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            return True
    return False
