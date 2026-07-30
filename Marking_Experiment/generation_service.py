"""Reusable task JSON generation for the Marking Experiment workflow."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from .marksheet_parser import MarksheetParser


SUPPORTED_INPUT_SUFFIXES = {
    ".txt",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".md",
    ".json",
    ".yaml",
    ".yml",
}


def infer_program(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return "excel"
    if suffix in {".html", ".htm"}:
        return "html"
    if suffix in {".accdb", ".mdb"}:
        return "access"
    return "word"


def classify_document_role(path: Path) -> str:
    name = path.name.lower()
    if re.search(r"\bmg\b|memo|marking guideline|rubric", name):
        return "memo"
    if "learner" in name or "learner data" in name:
        return "learner"
    if re.search(r"\bqp\b|question|paper|p1", name):
        return "question_paper"

    if path.suffix.lower() == ".docx":
        try:
            from docx import Document

            document = Document(path)
            headers = []
            for table in document.tables[:5]:
                if table.rows:
                    headers.append(" ".join(cell.text.lower() for cell in table.rows[0].cells))
            header_text = " ".join(headers)
            paragraph_text = " ".join(p.text.lower() for p in document.paragraphs[:10])

            if "criteria" in header_text and ("max mark" in header_text or "learner mark" in header_text):
                if "question" in header_text and "marks" in header_text:
                    return "combined"
                return "memo"
            if "question" in header_text and "marks" in header_text:
                return "question_paper"
            if "marking guideline" in paragraph_text:
                return "memo"
        except Exception:
            pass

    return "unknown"


def read_input_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".csv", ".md", ".json", ".yaml", ".yml"}:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="cp1252")

    if suffix == ".docx":
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        align_names = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        }

        document = Document(path)
        parts: List[str] = []
        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            hints = []
            fmt = paragraph.paragraph_format
            align = align_names.get(paragraph.alignment)
            if align:
                hints.append(f"align={align}")
            if fmt.line_spacing is not None:
                hints.append(f"line_spacing={fmt.line_spacing}")
            if fmt.space_before is not None:
                try:
                    hints.append(f"space_before={fmt.space_before.pt}pt")
                except Exception:
                    pass
            if fmt.space_after is not None:
                try:
                    hints.append(f"space_after={fmt.space_after.pt}pt")
                except Exception:
                    pass
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.size:
                        try:
                            hints.append(f"font_size={run.font.size.pt}pt")
                        except Exception:
                            pass
                    if run.font.bold:
                        hints.append("bold")
                    if run.font.italic:
                        hints.append("italic")
                    if run.font.name:
                        hints.append(f"font={run.font.name}")
                    break
            style_name = paragraph.style.name if paragraph.style else ""
            hint_str = f" [{', '.join(hints)}]" if hints else ""
            style_str = f" (style: {style_name})" if style_name else ""
            parts.append(f"[P{idx}]{style_str} {text}{hint_str}")

        for table_idx, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                rows.append("\t".join(cells))
            parts.append(f"TABLE {table_idx}:\n" + "\n".join(rows))

        return "\n\n".join(parts)

    if suffix in {".xlsx", ".xls"}:
        from openpyxl import load_workbook

        workbook = load_workbook(path, data_only=True)
        worksheet_texts = []
        for sheet in workbook.worksheets:
            worksheet_texts.append(f"Sheet: {sheet.title}")
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join(str(cell) if cell is not None else "" for cell in row))
            worksheet_texts.append("\n".join(rows))
        return "\n\n".join(worksheet_texts)

    raise RuntimeError(
        f"Unsupported input file type '{suffix}'. Use .txt, .docx, .xlsx, .xls, .csv, .md, .json, .yaml or .yml."
    )


def generate_marking_task_json(
    question_paper_path: Path,
    memo_path: Path,
    task_name: str,
    model: Optional[str] = None,
    progress_cb: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    program = infer_program(question_paper_path)
    question_paper_text = read_input_text(question_paper_path)
    memo_text = read_input_text(memo_path)

    parser = MarksheetParser(model=model)
    heuristic_questions = parser._parse_memo_tables(memo_text, program)
    if heuristic_questions:
        heuristic_questions = parser._atomize_questions(heuristic_questions)
        heuristic_questions = filter_questions_by_question_paper(question_paper_text, heuristic_questions)

    warnings: List[str] = []
    llm_output = ""
    llm_normalized_output = ""
    llm_status = "not_used"
    llm_questions: List[Dict[str, Any]] = []

    try:
        prompt = parser._build_prompt(question_paper_text, memo_text, program, str(memo_path))
        llm_output = parser._call_ollama(prompt, progress_cb=progress_cb)
        llm_questions = parser._atomize_questions(parser._extract_json(llm_output))
        llm_questions = normalize_llm_questions(llm_questions)
        llm_questions = filter_questions_by_question_paper(question_paper_text, llm_questions)
        warnings.extend(parser._validate_questions(llm_questions))
        llm_task_definition = build_task_definition(
            task_name=task_name,
            program=program,
            source_suffix=question_paper_path.suffix,
            questions=llm_questions,
        )
        llm_normalized_output = json.dumps(llm_task_definition, indent=2, ensure_ascii=False)
        llm_status = "success"
    except Exception as exc:
        warnings.append(f"LLM refinement unavailable; using heuristic JSON only. LLM error: {exc}")
        llm_status = "failed"

    questions = _reconcile_questions(llm_questions, heuristic_questions, warnings) if llm_questions else heuristic_questions
    questions = filter_questions_by_question_paper(question_paper_text, questions)
    task_definition = build_task_definition(
        task_name=task_name,
        program=program,
        source_suffix=question_paper_path.suffix,
        questions=questions,
    )

    json.dumps(task_definition, ensure_ascii=False)
    return {
        "task_definition": task_definition,
        "warnings": warnings,
        "program": program,
        "llm_status": llm_status,
        "llm_output": llm_output,
        "llm_normalized_output": llm_normalized_output,
        "heuristic_count": len(heuristic_questions),
        "llm_count": len(llm_questions),
    }


def build_task_definition(
    task_name: str,
    program: str,
    source_suffix: str,
    questions: List[Dict[str, Any]],
) -> Dict[str, Any]:
    normalized_questions = [dict(question, marks=1) for question in questions if isinstance(question, dict)]
    return {
        "task_name": task_name or "Generated Marking Task",
        "program": program,
        "file": f"student_file{source_suffix.lower() or '.docx'}",
        "total_marks": sum(int(question.get("marks", 1)) for question in normalized_questions),
        "questions": normalized_questions,
    }


def normalize_llm_questions(questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normalized = []
    for question in questions:
        q = dict(question)
        domain = str(q.get("domain", "")).strip()
        check_type = str(q.get("type", "")).strip()
        expected = q.get("expected")
        target = q.get("target") if isinstance(q.get("target"), dict) else {}

        if isinstance(expected, dict):
            expected = _normalize_expected_dict(check_type, expected)

        if domain == "paragraph_formatting" and check_type == "indent_level":
            check_type = "left_indent"
            if isinstance(expected, dict) and expected.get("value") is not None:
                expected = expected["value"]

        if check_type == "alignment" and isinstance(expected, dict):
            expected = str(expected.get("alignment", "")).lower().replace("justified", "justify")

        if check_type == "line_spacing" and isinstance(expected, dict):
            expected = {
                "rule": expected.get("rule", "exact"),
                "value": expected.get("value", 0),
                "unit": expected.get("unit", "pt"),
            }

        if check_type == "space_before" and isinstance(expected, dict) and expected.get("value") is not None:
            expected = expected["value"]

        if check_type == "first_line_indent" and isinstance(expected, dict) and expected.get("value") is not None:
            expected = expected["value"]

        if check_type == "border" and isinstance(expected, dict):
            border_expected = {}
            if expected.get("width_pt") is not None:
                border_expected["width_pt"] = expected["width_pt"]
            if expected.get("color") is not None:
                border_expected["color"] = expected["color"]
            expected = border_expected or True

        if check_type == "shading" and isinstance(expected, dict):
            expected = {"color": expected.get("color", expected.get("shading_color", "any"))}

        if check_type == "drop_cap" and isinstance(expected, dict):
            expected = {"lines": int(expected.get("lines", expected.get("lines_to_drop", 3)))}

        if target.get("locator") == "after_text":
            target = {"locator": "contains_text", "value": target.get("value", "")}

        q["domain"] = domain
        q["type"] = check_type
        q["target"] = target
        q["expected"] = expected
        q["marks"] = 1
        normalized.append(q)
    return normalized


def _normalize_expected_dict(check_type: str, expected: Dict[str, Any]) -> Dict[str, Any]:
    out = dict(expected)

    def parse_unit_value(value: Any) -> tuple[Optional[float], Optional[str]]:
        match = re.search(r"(\d+(?:\.\d+)?)\s*(cm|pt|lines?)?", str(value).lower())
        if not match:
            return None, None
        return float(match.group(1)), (match.group(2) or None)

    for source_key in ("indent_level", "line_spacing", "space_before", "border_width"):
        if source_key in out:
            value, unit = parse_unit_value(out[source_key])
            if value is not None:
                out["value" if source_key != "border_width" else "width_pt"] = value
            if unit:
                out["unit"] = "lines" if unit.startswith("line") else unit

    if "border_color" in out:
        out["color"] = out["border_color"]
    if "shading_color" in out:
        out["color"] = out["shading_color"]
    if "lines_to_drop" in out:
        out["lines"] = out["lines_to_drop"]
    if check_type == "line_spacing":
        out.setdefault("rule", "exact")
        out.setdefault("unit", "pt")
    return out


def filter_questions_by_question_paper(question_paper_text: str, questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Drop memo-only checks that are not supported by the question paper.

    Memos sometimes contain reused setup criteria that are absent from the actual
    question paper. The heuristic stage should not mark learners on those.
    """
    qp = question_paper_text.lower()
    required_terms_by_type = {
        "paper_size": ("paper size", "page size", "document size", "a4"),
        "orientation": ("orientation", "portrait", "landscape"),
        "margins": ("margin",),
        "hyphenation": ("hyphenation",),
        "page_border": ("page border",),
        "watermark": ("watermark",),
        "header_text": ("header",),
        "header_alignment": ("header",),
        "footer_text": ("footer", "page number"),
        "footer_alignment": ("footer",),
    }

    filtered: List[Dict[str, Any]] = []
    for question in questions:
        check_type = str(question.get("type", "")).lower()
        required_terms = required_terms_by_type.get(check_type)
        if required_terms and not any(term in qp for term in required_terms):
            continue
        filtered.append(question)
    return filtered


def _question_key(question: Dict[str, Any]) -> str:
    return str(question.get("question_number", "")).strip().lower()


def _question_signature(question: Dict[str, Any]) -> str:
    comparable = {
        "domain": question.get("domain"),
        "type": question.get("type"),
        "target": question.get("target"),
        "expected": question.get("expected"),
    }
    return json.dumps(comparable, sort_keys=True, ensure_ascii=False, default=str)


def _reconcile_questions(
    llm_questions: List[Dict[str, Any]],
    heuristic_questions: List[Dict[str, Any]],
    warnings: List[str],
) -> List[Dict[str, Any]]:
    reconciled = [dict(question, marks=1) for question in llm_questions]
    by_number = {_question_key(question): question for question in reconciled if _question_key(question)}
    existing_signatures = {_question_signature(question) for question in reconciled}

    for heuristic in heuristic_questions:
        heuristic = dict(heuristic, marks=1)
        key = _question_key(heuristic)
        signature = _question_signature(heuristic)

        if signature in existing_signatures:
            continue

        if key and key in by_number:
            warnings.append(
                f"Heuristic and LLM disagreed on question {heuristic.get('question_number')}; kept LLM version."
            )
            continue

        reconciled.append(heuristic)
        existing_signatures.add(signature)
        if key:
            by_number[key] = heuristic
        warnings.append(
            f"Added heuristic check missing from LLM output: {heuristic.get('question_number', 'unknown')}."
        )

    return reconciled
